import os
import sys
import csv
import json
import datetime
import subprocess
import tempfile
import uuid
from dataclasses import dataclass, field, asdict
from typing import List, Optional, Dict

import tkinter as tk
from tkinter import ttk, filedialog, messagebox

# --- Packaged app helpers (exe-safe paths) ---
def resource_path(*parts):
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        base = sys._MEIPASS
    else:
        base = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base, *parts)

def app_writable_dir():
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.getcwd()

# ---------------- Optional export deps ----------------
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Paragraph, SimpleDocTemplate, Spacer, Image as RLImage,
        ListFlowable, ListItem, Table, TableStyle
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
except Exception:
    SimpleDocTemplate = None

try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:
    docx2pdf_convert = None

# --- Optional Pillow for image preview/edit ---
try:
    from PIL import Image, ImageTk
except Exception:
    Image = None
    ImageTk = None

# ------------------ Helpers & constants ------------------

CHEM_SYMBOL_SETS = [
    ("→", "Right arrow"), ("←", "Left arrow"), ("⇌", "Eqm arrow"),
    ("⇄", "Eqm (alt)"), ("↔", "Reversible"), ("⇑", "Up"), ("⇓", "Down"),
    ("α", "alpha"), ("β", "beta"), ("γ", "gamma"), ("δ", "delta"), ("ε", "epsilon"),
    ("λ", "lambda"), ("μ", "mu"), ("π", "pi"), ("σ", "sigma"), ("φ", "phi"), ("ω", "omega"),
    ("⁺", "sup +"), ("⁻", "sup -"), ("¹", "sup 1"), ("²", "sup 2"), ("³", "sup 3"),
    ("₊", "sub +"), ("₋", "sub -"), ("₁", "sub 1"), ("₂", "sub 2"), ("₃", "sub 3"),
    ("°", "degree"), ("±", "plus/minus"), ("∙", "dot"), ("·", "middot"), ("®", "reg"),
]
CHEM_SYMBOL_OPTIONS = [f"{sym}  —  {desc}" for sym, desc in CHEM_SYMBOL_SETS]

DOCX_TABLE_STYLES = [
    "Table Grid",
    "Light List",
    "Light List Accent 1",
    "Light Grid",
    "Light Grid Accent 1",
    "Medium Grid 1",
    "Medium Grid 1 Accent 1",
    "Medium Shading 1",
    "Medium Shading 1 Accent 1",
]

TRIAL_LAYOUTS = ["Even columns", "Reasons wide", "Compact"]

def safe_split_lines(text: str) -> List[str]:
    return [ln.strip() for ln in (text or "").splitlines() if ln.strip()]

def ensure_dir(path: str):
    os.makedirs(os.path.dirname(path), exist_ok=True)

def parse_table_text(text: str) -> List[List[str]]:
    txt = (text or "").strip("\n")
    if not txt:
        return []
    try:
        sniff = csv.Sniffer().sniff(txt, delimiters="\t,;")
        rows = list(csv.reader(txt.splitlines(), sniff))
        if rows and any(len(r) > 1 for r in rows):
            return [[c.strip() for c in r] for r in rows]
    except Exception:
        pass
    if "\t" in txt:
        return [[c.strip() for c in ln.split("\t")] for ln in txt.splitlines()]
    import re
    rows = []
    for ln in txt.splitlines():
        parts = [p for p in re.split(r"\s{2,}", ln.strip()) if p != ""]
        rows.append(parts if parts else [ln.strip()])
    return rows

# --------------------------- Data models ---------------------------

@dataclass
class TrialRow:
    number: str
    issue: str
    reasons: str

@dataclass
class ResultItem:
    title: str
    kind: str  # 'text' | 'table' | 'image'
    content: str = ""
    image_path: Optional[str] = None
    images: List[str] = field(default_factory=list)
    caption: str = ""
    table_data: List[List[str]] = field(default_factory=list)
    table_style: str = "Table Grid"

@dataclass
class EmailEntry:
    date: str
    customer: str
    correspondence: str

@dataclass
class ReportModel:
    # 1. General
    project_title: str = ""
    start_date: str = ""
    report_date: str = ""
    assigned_by: str = ""
    bin_no: str = ""
    researcher_name: str = ""
    total_hours: str = ""

    # 2. Technical
    plain_summary: str = ""
    objectives: List[str] = field(default_factory=list)
    methods_raw_materials: List[str] = field(default_factory=list)
    methods_instruments: List[str] = field(default_factory=list)
    methods_procedure: List[str] = field(default_factory=list)
    trial_history: List[TrialRow] = field(default_factory=list)
    trial_layout: str = TRIAL_LAYOUTS[0]
    trial_docx_style: str = DOCX_TABLE_STYLES[0]
    results: List[ResultItem] = field(default_factory=list)
    conclusion: List[str] = field(default_factory=list)
    miscellaneous: str = ""
    references: List[str] = field(default_factory=list)

    # 3. Regulatory
    regulations: str = ""
    label_req: str = ""
    certification_req: str = ""

    # 4. Scale Up
    manuf_order_steps: List[str] = field(default_factory=list)  # 4.1.
    formulation_risk_text: str = ""       # 4.2.
    hazards_text: str = ""                # 4.3.
    equipment_text: str = ""              # 4.4.
    capex_text: str = ""                  # 4.5.
    safety_assess_text: str = ""          # 4.6.

    # 5. Quality
    raw_material_sourcing: str = ""
    lims_setup: str = ""
    stability_testing: str = ""
    packaging_compatibility: str = ""

    # 6. Commercial (expanded 6.1.–6.12.)
    c_obj_problem: str = ""  # 6.1.
    smart_goals: Dict[str, str] = field(default_factory=lambda: {"S":"","M":"","A":"","R":"","T":""})  # 6.2.
    c_specs: str = ""  # 6.3.
    c_expected_volume: str = ""  # 6.4.
    c_packaging_req: str = ""  # 6.5.
    c_raw_material_prefs: str = ""  # 6.6.
    c_sample_needed: str = ""  # 6.7.
    c_opportunity_timeline: str = ""  # 6.8.
    target_application: str = ""  # 6.9.
    customer_feedback: str = ""  # 6.10.
    tds_development: str = ""  # 6.11.
    email_correspondence: List[EmailEntry] = field(default_factory=list)  # 6.12.

# ----------------------- Tooltip / “ii” UI bits -----------------------
class Tooltip:
    def __init__(self, widget, text, wrap=300):
        self.widget, self.text, self.wrap = widget, text, wrap
        self.tipwin = None
        widget.bind("<Enter>", self.show)
        widget.bind("<Leave>", self.hide)
    def show(self, _e=None):
        if self.tipwin or not self.text: return
        x = self.widget.winfo_rootx() + 10
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 6
        self.tipwin = tw = tk.Toplevel(self.widget); tw.wm_overrideredirect(True); tw.wm_geometry(f"+{x}+{y}")
        frame = ttk.Frame(tw, padding=6, style="Tooltip.TFrame"); frame.pack()
        ttk.Label(frame, text=self.text, justify="left", wraplength=self.wrap).pack()
    def hide(self, _e=None):
        if self.tipwin:
            self.tipwin.destroy(); self.tipwin = None

def info_icon(parent, text: str):
    c = tk.Canvas(parent, width=16, height=16, highlightthickness=0)
    c.create_oval(1, 1, 15, 15, fill="#0b62a9", outline="#0b62a9")
    c.create_text(8, 8, text="i", fill="white", font=("Segoe UI", 10, "bold"))
    c.pack(side=tk.LEFT, padx=(6, 0))
    Tooltip(c, text)
    return c

class ChemDropdown(ttk.Frame):
    def __init__(self, master, get_target_widget_callable, **kwargs):
        super().__init__(master, **kwargs)
        ttk.Label(self, text="Chem symbol:").pack(side=tk.LEFT, padx=(0, 4))
        self.get_target = get_target_widget_callable
        self.var = tk.StringVar()
        self.combo = ttk.Combobox(self, values=CHEM_SYMBOL_OPTIONS, textvariable=self.var,
                                  width=24, state="readonly", takefocus=False)
        self.combo.pack(side=tk.LEFT)
        Tooltip(self.combo, "Pick a symbol; it inserts into the last field you edited.")
        self.combo.bind("<<ComboboxSelected>>", lambda _e: self._insert())
        self.combo.bind("<FocusIn>", lambda _e: self.after(1, self._restore_focus))
        self.combo.bind("<Return>", lambda _e: self._insert())
    def _restore_focus(self):
        w = self.get_target()
        if w and w.winfo_exists():
            try: w.focus_set()
            except Exception: pass
    def _insert(self):
        choice = self.var.get().strip()
        if not choice: return
        sym = choice.split("—", 1)[0].strip()
        w = self.get_target()
        if not (w and w.winfo_exists()):
            w = self.focus_get()
            if not (w and str(w.winfo_class()) in ("Text", "Entry", "TEntry")):
                w = None
        if not w:
            self.var.set(""); return
        try:
            w.insert("insert", sym); w.focus_set()
        except Exception:
            pass
        self.var.set("")

# ------------------------------ Image Editor ------------------------------
class ImageManager(ttk.Frame):
    def __init__(self, master, initial_files: Optional[List[str]] = None, preview_size=(520, 360)):
        super().__init__(master)
        self.preview_size = preview_size
        self.files: List[str] = list(initial_files or [])
        self._images_cache = {}
        self._orig_cache = {}
        self._tk_img = None

        left = ttk.Frame(self); left.pack(side=tk.LEFT, fill="y", padx=(0, 8))
        ttk.Label(left, text="Selected Images:").pack(anchor="w")
        self.lst = tk.Listbox(left, height=18, selectmode="extended"); self.lst.pack(fill="y", expand=False)
        btns1 = ttk.Frame(left); btns1.pack(fill="x", pady=4)
        ttk.Button(btns1, text="Add…", command=self._add_files).pack(side=tk.LEFT)
        ttk.Button(btns1, text="Remove", command=self._remove_selected).pack(side=tk.LEFT, padx=6)
        ttk.Button(btns1, text="Up", command=lambda: self._move(-1)).pack(side=tk.LEFT, padx=6)
        ttk.Button(btns1, text="Down", command=lambda: self._move(1)).pack(side=tk.LEFT)

        right = ttk.Frame(self); right.pack(side=tk.LEFT, fill="both", expand=True)
        self.preview = tk.Label(right, relief="sunken", width=64, height=20, anchor="center")
        self.preview.pack(fill="both", expand=True)
        editbar = ttk.Frame(right); editbar.pack(fill="x", pady=(6, 0))
        self._supports_edit = (Image is not None and ImageTk is not None)
        if self._supports_edit:
            ttk.Button(editbar, text="Rotate ⟲", command=self._rotate_left).pack(side=tk.LEFT)
            ttk.Button(editbar, text="Rotate ⟳", command=self._rotate_right).pack(side=tk.LEFT, padx=6)
            ttk.Button(editbar, text="Flip ↔", command=self._flip_h).pack(side=tk.LEFT, padx=6)
            ttk.Button(editbar, text="Flip ↕", command=self._flip_v).pack(side=tk.LEFT, padx=6)
            ttk.Button(editbar, text="Reset", command=self._reset).pack(side=tk.LEFT, padx=12)
        else:
            ttk.Label(editbar, text="Pillow not installed — preview/edit disabled", foreground="#a00").pack(side=tk.LEFT)

        for f in self.files:
            self.lst.insert("end", f)
        self.lst.bind("<<ListboxSelect>>", lambda _e: self._render_preview())

    def get_files(self) -> List[str]:
        return list(self.files)

    def _add_files(self):
        paths = filedialog.askopenfilenames(
            title="Select image(s)",
            filetypes=[("Images", "*.png;*.jpg;*.jpeg;*.bmp;*.tif;*.tiff;*.gif")]
        )
        if not paths: return
        for p in paths:
            if p not in self.files:
                self.files.append(p); self.lst.insert("end", p)
        if self.files:
            self.lst.selection_clear(0, "end"); self.lst.selection_set("end"); self.lst.activate("end")
            self._render_preview()

    def _remove_selected(self):
        sel = list(self.lst.curselection())
        if not sel: return
        sel.reverse()
        for i in sel:
            path = self.lst.get(i); self.lst.delete(i)
            if path in self.files: self.files.remove(path)
            self._images_cache.pop(path, None); self._orig_cache.pop(path, None)
        self._render_preview()

    def _move(self, delta: int):
        i = self.lst.curselection()
        if not i: return
        i = i[0]; j = i + delta
        if j < 0 or j >= self.lst.size(): return
        a = self.lst.get(i); b = self.lst.get(j)
        self.lst.delete(j); self.lst.insert(j, a)
        self.lst.delete(i); self.lst.insert(i, b)
        self.lst.selection_clear(0, "end"); self.lst.selection_set(j); self.lst.activate(j)
        self.files[i], self.files[j] = self.files[j], self.files[i]
        self._render_preview()

    def _load_image(self, path: str):
        if not self._supports_edit or not os.path.isfile(path): return None
        if path not in self._orig_cache:
            try: self._orig_cache[path] = Image.open(path).convert("RGBA")
            except Exception: return None
        if path not in self._images_cache:
            self._images_cache[path] = self._orig_cache[path].copy()
        return self._images_cache[path]

    def _render_preview(self):
        if not self._supports_edit:
            self.preview.config(text="(Preview unavailable)"); return
        if not self.files:
            self.preview.config(image="", text="No image selected"); return
        idxs = self.lst.curselection()
        idx = idxs[0] if idxs else min(len(self.files)-1, 0)
        path = self.files[idx]
        im = self._load_image(path)
        if im is None:
            self.preview.config(text="Could not load image"); return
        max_w, max_h = self.preview_size
        w, h = im.size
        scale = min(max_w / max(1, w), max_h / max(1, h))
        size = (max(1, int(w*scale)), max(1, int(h*scale)))
        disp = im.resize(size, Image.LANCZOS)
        self._tk_img = ImageTk.PhotoImage(disp)
        self.preview.config(image=self._tk_img, text="")

    def _selected_path(self) -> Optional[str]:
        if not self.files: return None
        sel = self.lst.curselection()
        if not sel: return self.files[-1]
        return self.files[sel[0]]

    def _rotate_left(self):
        p = self._selected_path(); 
        if not p: return
        im = self._load_image(p)
        if im is None: return
        self._images_cache[p] = im.rotate(90, expand=True); self._render_preview()

    def _rotate_right(self):
        p = self._selected_path(); 
        if not p: return
        im = self._load_image(p)
        if im is None: return
        self._images_cache[p] = im.rotate(-90, expand=True); self._render_preview()

    def _flip_h(self):
        p = self._selected_path(); 
        if not p: return
        im = self._load_image(p)
        if im is None: return
        self._images_cache[p] = im.transpose(Image.FLIP_LEFT_RIGHT); self._render_preview()

    def _flip_v(self):
        p = self._selected_path(); 
        if not p: return
        im = self._load_image(p)
        if im is None: return
        self._images_cache[p] = im.transpose(Image.FLIP_TOP_BOTTOM); self._render_preview()

    def _reset(self):
        p = self._selected_path()
        if not p: return
        if p in self._orig_cache:
            self._images_cache[p] = self._orig_cache[p].copy(); self._render_preview()

    def save_edited_to_temp(self) -> List[str]:
        out_paths: List[str] = []
        cache_root = os.path.join(tempfile.gettempdir(), "rdr_imgcache")
        os.makedirs(cache_root, exist_ok=True)
        for path in self.files:
            base = os.path.basename(path)
            name, ext = os.path.splitext(base)
            ext = ext if ext else ".png"
            out = os.path.join(cache_root, f"{name}_{uuid.uuid4().hex[:8]}{ext}")
            if self._supports_edit and path in self._images_cache:
                try:
                    fmt = None
                    if ext.lower() in [".jpg", ".jpeg"]: fmt = "JPEG"
                    elif ext.lower() == ".png": fmt = "PNG"
                    elif ext.lower() in [".tif", ".tiff"]: fmt = "TIFF"
                    self._images_cache[path].save(out, format=fmt)
                except Exception:
                    try:
                        from shutil import copy2; copy2(path, out)
                    except Exception:
                        continue
            else:
                try:
                    from shutil import copy2; copy2(path, out)
                except Exception:
                    continue
            out_paths.append(out)
        return out_paths

# ------------------------------ Main App ------------------------------
class ReportApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("R&D Final Report Builder")
        self.geometry("1200x900")
        self.report = ReportModel()
        self._last_focused_widget: Optional[tk.Widget] = None

        self.current_file_path: Optional[str] = None
        self._include_state_keys = [
            "sec_general", "t_plain", "t_objectives", "t_methods", "t_rm", "t_ins", "t_proc",
            "t_trial", "t_results", "t_conc", "t_misc", "t_refs",
            "sec_reg", "sec_scale", "sec_quality", "sec_commercial"
        ]

        self.header_path_var = tk.StringVar(value=resource_path("assets", "Header.png"))
        self.footer_path_var = tk.StringVar(value=resource_path("assets", "Footer.png"))
        self.logo_path_var   = tk.StringVar(value=resource_path("assets", "Logo_Only.png"))

        self._ui_logo_small = None
        self._ui_logo_large = None

        self._build_styles()
        self._prepare_ui_logo()

        self._tab_roots: List[tk.Widget] = []

        self._make_menubar()
        self._branding_bar()
        self._build_ui()
        self._wire_focus_tracking()

        self.bind_all("<Control-n>", lambda e: self._file_new())
        self.bind_all("<Control-o>", lambda e: self._file_open())
        self.bind_all("<Control-s>", lambda e: self._file_save())
        self.bind_all("<Control-S>", lambda e: self._file_save_as())

    def _build_styles(self):
        style = ttk.Style(self)
        try:
            if sys.platform == "darwin":
                style.theme_use("clam")
        except Exception:
            pass
        style.configure("Tooltip.TFrame", background="#ffffe0", relief="solid", borderwidth=1)
        style.configure("h1.TLabel", font=("Segoe UI", 14, "bold"))

    def _prepare_ui_logo(self):
        path = self.logo_path_var.get()
        self._ui_logo_small = None; self._ui_logo_large = None
        if not (path and os.path.isfile(path)): return
        try:
            small = tk.PhotoImage(file=path)
            factor = max(1, small.width() // 10)
            self._ui_logo_small = small.subsample(factor, factor)
            try: self.iconphoto(True, small)
            except Exception: pass
        except Exception:
            pass
        if Image and ImageTk:
            try:
                im = Image.open(path).convert("RGBA")
                target_w = 280
                if im.width > target_w:
                    ratio = target_w / float(im.width)
                    im = im.resize((target_w, max(1, int(im.height * ratio))), Image.LANCZOS)
                r,g,b,a = im.split()
                a = a.point(lambda px: int(px * 0.16))
                im = Image.merge("RGBA", (r,g,b,a))
                self._ui_logo_large = ImageTk.PhotoImage(im)
            except Exception:
                self._ui_logo_large = None
        else:
            try:
                base = tk.PhotoImage(file=path)
                factor = max(1, base.width() // 220)
                self._ui_logo_large = base.subsample(factor, factor)
            except Exception:
                self._ui_logo_large = None

    def _apply_ui_watermark(self, container):
        if not self._ui_logo_large: return
        lbl = tk.Label(container, image=self._ui_logo_large, borderwidth=0, highlightthickness=0)
        lbl.place(relx=0.5, rely=0.6, anchor="center")
        lbl.lower()
        if not hasattr(container, "_wm_labels"):
            container._wm_labels = []
        container._wm_labels.append(lbl)

    def _refresh_all_watermarks(self):
        for root in getattr(self, "_tab_roots", []):
            if hasattr(root, "_wm_labels"):
                for w in root._wm_labels:
                    try: w.destroy()
                    except Exception: pass
                root._wm_labels.clear()
            self._apply_ui_watermark(root)

    def _watch_widget(self, w: tk.Widget):
        w.bind("<FocusIn>", lambda e: setattr(self, "_last_focused_widget", w), add="+")

    def _make_menubar(self):
        menubar = tk.Menu(self)
        filemenu = tk.Menu(menubar, tearoff=0)
        filemenu.add_command(label="New", accelerator="Ctrl+N", command=self._file_new)
        filemenu.add_separator()
        filemenu.add_command(label="Open…", accelerator="Ctrl+O", command=self._file_open)
        filemenu.add_separator()
        filemenu.add_command(label="Save", accelerator="Ctrl+S", command=self._file_save)
        filemenu.add_command(label="Save As…", accelerator="Ctrl+Shift+S", command=self._file_save_as)
        menubar.add_cascade(label="File", menu=filemenu)
        self.config(menu=menubar)

    # File menu actions
    def _file_new(self):
        try:
            if getattr(sys, "frozen", False):
                subprocess.Popen([sys.executable])
            else:
                subprocess.Popen([sys.executable, os.path.abspath(__file__)])
        except Exception as e:
            messagebox.showerror("New Project", f"Could not launch new window:\n{e}")

    def _file_open(self):
        path = filedialog.askopenfilename(
            title="Open Project",
            filetypes=[("R&D Report Project", "*.rdrproj"), ("JSON", "*.json"), ("All files", "*.*")]
        )
        if not path: return
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            self._load_from_dict(data)
            self.current_file_path = path
            self._refresh_title()
            messagebox.showinfo("Open", f"Loaded: {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Open Failed", f"Could not open file:\n{e}")

    def _file_save(self):
        if not self.current_file_path:
            return self._file_save_as()
        try:
            data = self._to_dict()
            with open(self.current_file_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2)
            messagebox.showinfo("Save", f"Saved: {os.path.basename(self.current_file_path)}")
        except Exception as e:
            messagebox.showerror("Save Failed", f"Could not save file:\n{e}")

    def _file_save_as(self):
        default_name = (self.ent_title.get().strip() or "project") + ".rdrproj"
        path = filedialog.asksaveasfilename(
            title="Save Project As",
            defaultextension=".rdrproj",
            initialfile=default_name,
            filetypes=[("R&D Report Project", "*.rdrproj"), ("JSON", "*.json"), ("All files", "*.*")]
        )
        if not path: return
        try:
            data = self._to_dict()
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2)
            self.current_file_path = path
            self._refresh_title()
            messagebox.showinfo("Save As", f"Saved: {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Save As Failed", f"Could not save file:\n{e}")

    def _refresh_title(self):
        base = "R&D Final Report Builder"
        if self.current_file_path:
            base += f" — {os.path.basename(self.current_file_path)}"
        self.title(base)

    # persistence
    def _to_dict(self) -> dict:
        r = self._collect()
        include_state = {k: bool(self.include[k].get()) for k in self._include_state_keys}
        data = {
            "report_model": {
                **asdict(r),
                "trial_history": [asdict(tr) for tr in r.trial_history],
                "results": [asdict(res) for res in r.results],
                "email_correspondence": [asdict(e) for e in r.email_correspondence],
            },
            "branding": {
                "header_path": self.header_path_var.get(),
                "footer_path": self.footer_path_var.get(),
                "logo_path": self.logo_path_var.get(),
            },
            "export": {
                "out_dir": self.out_dir_var.get(),
                "format": self.export_fmt.get(),
            },
            "include": include_state,
        }
        return data

    def _load_from_dict(self, data: dict):
        try:
            rm = data.get("report_model", {})
            results_in = []
            for ri in rm.get("results", []):
                if "images" not in ri:
                    ri["images"] = []
                    if ri.get("image_path"):
                        ri["images"] = [ri["image_path"]]
                results_in.append(ri)

            emails_in = []
            for e in rm.get("email_correspondence", []):
                emails_in.append(e)

            # Back-compat: if old 'manuf_order_text' exists, split it into steps
            mo_steps = list(rm.get("manuf_order_steps", []))
            if not mo_steps and rm.get("manuf_order_text", ""):
                mo_steps = safe_split_lines(rm.get("manuf_order_text", ""))

            # ---------------------- FIX IS HERE ----------------------
            r = ReportModel(
                project_title=rm.get("project_title", ""),
                start_date=rm.get("start_date", ""),
                report_date=rm.get("report_date", ""),
                assigned_by=rm.get("assigned_by", ""),
                bin_no=rm.get("bin_no", ""),
                researcher_name=rm.get("researcher_name", ""),
                total_hours=rm.get("total_hours", ""),
                plain_summary=rm.get("plain_summary", ""),
                objectives=list(rm.get("objectives", [])),  # <- use rm, not r
                methods_raw_materials=list(rm.get("methods_raw_materials", [])),
                methods_instruments=list(rm.get("methods_instruments", [])),
                methods_procedure=list(rm.get("methods_procedure", [])),
                trial_history=[TrialRow(**th) for th in rm.get("trial_history", [])],
                trial_layout=rm.get("trial_layout", TRIAL_LAYOUTS[0]),
                trial_docx_style=rm.get("trial_docx_style", DOCX_TABLE_STYLES[0]),
                results=[ResultItem(**ri) for ri in results_in],
                conclusion=list(rm.get("conclusion", [])),
                miscellaneous=rm.get("miscellaneous", ""),
                references=list(rm.get("references", [])),

                regulations=rm.get("regulations", ""),
                label_req=rm.get("label_req", ""),
                certification_req=rm.get("certification_req", ""),

                # Section 4 (with new 4.1 list)
                manuf_order_steps=mo_steps,
                formulation_risk_text=rm.get("formulation_risk_text",""),
                hazards_text=rm.get("hazards_text",""),
                equipment_text=rm.get("equipment_text",""),
                capex_text=rm.get("capex_text",""),
                safety_assess_text=rm.get("safety_assess_text",""),

                raw_material_sourcing=rm.get("raw_material_sourcing", ""),
                lims_setup=rm.get("lims_setup", ""),
                stability_testing=rm.get("stability_testing", ""),
                packaging_compatibility=rm.get("packaging_compatibility", ""),

                c_obj_problem=rm.get("c_obj_problem", ""),
                smart_goals=rm.get("smart_goals", {"S":"","M":"","A":"","R":"","T":""}),
                c_specs=rm.get("c_specs",""),
                c_expected_volume=rm.get("c_expected_volume",""),
                c_packaging_req=rm.get("c_packaging_req",""),
                c_raw_material_prefs=rm.get("c_raw_material_prefs",""),
                c_sample_needed=rm.get("c_sample_needed",""),
                c_opportunity_timeline=rm.get("c_opportunity_timeline",""),
                target_application=rm.get("target_application",""),
                customer_feedback=rm.get("customer_feedback",""),
                tds_development=rm.get("tds_development",""),
                email_correspondence=[EmailEntry(**e) for e in emails_in],
            )
            # ---------------------------------------------------------

            self._populate_from_model(r)

            b = data.get("branding", {})
            self.header_path_var.set(b.get("header_path", self.header_path_var.get()))
            self.footer_path_var.set(b.get("footer_path", self.footer_path_var.get()))
            self.logo_path_var.set(b.get("logo_path", self.logo_path_var.get()))
            self._prepare_ui_logo(); self._refresh_all_watermarks()

            e = data.get("export", {})
            self.out_dir_var.set(e.get("out_dir", self.out_dir_var.get()))
            self.export_fmt.set(e.get("format", self.export_fmt.get()))

            inc = data.get("include", {})
            for k in self._include_state_keys:
                if k in inc: self.include[k].set(bool(inc[k]))

            self.trial_layout_var.set(r.trial_layout)
            self.trial_style_var.set(r.trial_docx_style)
            self._update_trial_preview()

        except Exception as e:
            messagebox.showerror("Load Failed", f"Could not load project data:\n{e}")

    def _populate_from_model(self, r: ReportModel):
        # General
        self.ent_title.delete(0, "end"); self.ent_title.insert(0, r.project_title)
        self.start_date_var.set(r.start_date); self.report_date_var.set(r.report_date)
        self.assigned_by_var.set(r.assigned_by); self.bin_no_var.set(r.bin_no)
        self.researcher_var.set(r.researcher_name); self.total_hours_var.set(r.total_hours)

        # Technical
        self.txt_plain.delete("1.0", "end"); self.txt_plain.insert("1.0", r.plain_summary)

        for child in list(self._obj_container.winfo_children()): child.destroy()
        self.obj_vars = []
        if r.objectives:
            for obj in r.objectives:
                self._add_objective_row(); self.obj_vars[-1].set(obj)
        else:
            self._add_objective_row()

        self.txt_rm.delete("1.0","end"); self.txt_rm.insert("1.0","\n".join(r.methods_raw_materials))
        self.txt_ins.delete("1.0","end"); self.txt_ins.insert("1.0","\n".join(r.methods_instruments))
        self.txt_proc.delete("1.0","end"); self.txt_proc.insert("1.0","\n".join(r.methods_procedure))

        for iid in self.trial_tree.get_children(""):
            self.trial_tree.delete(iid)
        for tr in r.trial_history:
            self.trial_tree.insert("", "end", values=(tr.number, tr.issue, tr.reasons))

        self._results = list(r.results)
        self._rebuild_results_list()

        self.txt_conclusion.delete("1.0","end"); self.txt_conclusion.insert("1.0","\n".join(r.conclusion))
        self.txt_misc.delete("1.0","end"); self.txt_misc.insert("1.0", r.miscellaneous)
        self.txt_refs.delete("1.0","end"); self.txt_refs.insert("1.0","\n".join(r.references))

        # Regulatory
        self.txt_regs.delete("1.0","end"); self.txt_regs.insert("1.0", r.regulations)
        self.txt_label_req.delete("1.0","end"); self.txt_label_req.insert("1.0", r.label_req)
        self.txt_cert_req.delete("1.0","end"); self.txt_cert_req.insert("1.0", r.certification_req)

        # Scale Up (4.1.–4.6.)
        for child in list(self._mo_container.winfo_children()): child.destroy()
        self.mo_vars = []
        if r.manuf_order_steps:
            for step in r.manuf_order_steps:
                self._add_mo_row(); self.mo_vars[-1].set(step)
        else:
            self._add_mo_row()

        self.txt_42.delete("1.0","end"); self.txt_42.insert("1.0", r.formulation_risk_text)
        self.txt_43.delete("1.0","end"); self.txt_43.insert("1.0", r.hazards_text)
        self.txt_44.delete("1.0","end"); self.txt_44.insert("1.0", r.equipment_text)
        self.txt_45.delete("1.0","end"); self.txt_45.insert("1.0", r.capex_text)
        self.txt_46.delete("1.0","end"); self.txt_46.insert("1.0", r.safety_assess_text)

        # Quality
        self.txt_q_sourcing.delete("1.0","end"); self.txt_q_sourcing.insert("1.0", r.raw_material_sourcing)
        self.txt_q_lims.delete("1.0","end"); self.txt_q_lims.insert("1.0", r.lims_setup)
        self.txt_q_stability.delete("1.0","end"); self.txt_q_stability.insert("1.0", r.stability_testing)
        self.txt_q_pack.delete("1.0","end"); self.txt_q_pack.insert("1.0", r.packaging_compatibility)

        # Commercial (6.x.)
        self.txt_c_obj_problem.delete("1.0","end"); self.txt_c_obj_problem.insert("1.0", r.c_obj_problem)
        for k in ["S","M","A","R","T"]:
            self.smart_vars[k].set(r.smart_goals.get(k,""))
        self.txt_c_specs.delete("1.0","end"); self.txt_c_specs.insert("1.0", r.c_specs)
        self.txt_c_expected_volume.delete("1.0","end"); self.txt_c_expected_volume.insert("1.0", r.c_expected_volume)
        self.txt_c_packaging_req.delete("1.0","end"); self.txt_c_packaging_req.insert("1.0", r.c_packaging_req)
        self.txt_c_raw_material_prefs.delete("1.0","end"); self.txt_c_raw_material_prefs.insert("1.0", r.c_raw_material_prefs)
        self.txt_c_sample_needed.delete("1.0","end"); self.txt_c_sample_needed.insert("1.0", r.c_sample_needed)
        self.txt_c_opportunity_timeline.delete("1.0","end"); self.txt_c_opportunity_timeline.insert("1.0", r.c_opportunity_timeline)
        self.txt_c_target.delete("1.0","end"); self.txt_c_target.insert("1.0", r.target_application)
        self.txt_c_feedback.delete("1.0","end"); self.txt_c_feedback.insert("1.0", r.customer_feedback)
        self.txt_c_tds.delete("1.0","end"); self.txt_c_tds.insert("1.0", r.tds_development)

        # Email table
        for iid in self.email_tree.get_children(""):
            self.email_tree.delete(iid)
        for e in r.email_correspondence:
            self.email_tree.insert("", "end", values=(e.date, e.customer, e.correspondence))

    # branding bar
    def _branding_bar(self):
        bar = ttk.Frame(self, padding=(8, 6))
        bar.pack(fill="x")
        ttk.Label(bar, text="Thatcher Company — R&D Final Report Builder",
                  font=("Segoe UI", 12, "bold")).pack(side=tk.LEFT)
        if self._ui_logo_small:
            ttk.Label(bar, image=self._ui_logo_small).pack(side=tk.RIGHT, padx=6)
        def choose_logo():
            p = filedialog.askopenfilename(title="Select company logo", filetypes=[("Images", "*.png;*.jpg;*.jpeg;*.gif")])
            if not p: return
            self.logo_path_var.set(p); self._prepare_ui_logo()
            for child in list(bar.winfo_children()):
                # remove old image label(s)
                if isinstance(child, ttk.Label) and getattr(child, "image", None):
                    child.destroy()
            if self._ui_logo_small:
                ttk.Label(bar, image=self._ui_logo_small).pack(side=tk.RIGHT, padx=6)
            self._refresh_all_watermarks()
        ttk.Button(bar, text="Change Logo…", command=choose_logo).pack(side=tk.RIGHT, padx=6)

    def _wire_focus_tracking(self):
        def remember_focus(event):
            w = event.widget
            if str(w.winfo_class()) in ("Text", "Entry", "TEntry"):
                self._last_focused_widget = w
        self.bind_all("<FocusIn>", remember_focus, add="+")
    def _target_widget(self):
        return self._last_focused_widget

    # ---------- UI layout ----------
    def _build_ui(self):
        nb = ttk.Notebook(self); nb.pack(fill="both", expand=True)
        self._tab_general(nb)
        self._tab_technical(nb)
        self._tab_regulatory(nb)
        self._tab_scaleup(nb)
        self._tab_quality(nb)
        self._tab_commercial(nb)
        self._tab_export(nb)

    # ---- reusable widgets
    def labeled_entry(self, parent, label, textvariable=None, width=60, ii_text: Optional[str] = None):
        row = ttk.Frame(parent); row.pack(fill="x", pady=4)
        ttk.Label(row, text=label, width=24, anchor="w").pack(side=tk.LEFT)
        if ii_text: info_icon(row, ii_text)
        var = textvariable or tk.StringVar()
        ent = ttk.Entry(row, textvariable=var, width=width); ent.pack(side=tk.LEFT, fill="x", expand=True)
        self._watch_widget(ent)
        return var, ent

    def labeled_text(self, parent, label, height=6, ii_text: Optional[str] = None):
        frame = ttk.Frame(parent); frame.pack(fill="both", pady=4, expand=True)
        top = ttk.Frame(frame); top.pack(fill="x")
        ttk.Label(top, text=label, width=24, anchor="w").pack(side=tk.LEFT)
        if ii_text: info_icon(top, ii_text)
        txt = tk.Text(frame, height=height, wrap="word")
        txt.pack(fill="both", expand=True)
        self._watch_widget(txt)
        return txt

    def plain_text_area(self, parent, height=6, ii_text: Optional[str] = None):
        """Text area without a 'Details' label."""
        frame = ttk.Frame(parent); frame.pack(fill="both", pady=4, expand=True)
        if ii_text:
            top = ttk.Frame(frame); top.pack(fill="x")
            info_icon(top, ii_text)
        txt = tk.Text(frame, height=height, wrap="word")
        txt.pack(fill="both", expand=True)
        self._watch_widget(txt)
        return txt

    # ---------- Tabs
    def _tab_general(self, notebook):
        tab = ttk.Frame(notebook, padding=10)
        notebook.add(tab, text="1. General Information")
        self._tab_roots.append(tab); self._apply_ui_watermark(tab)
        _, self.ent_title = self.labeled_entry(tab, "Project Title:")
        self.start_date_var, _ = self.labeled_entry(tab, "Start Date (YYYY-MM-DD):")
        self.report_date_var, _ = self.labeled_entry(tab, "Report Date (YYYY-MM-DD):")
        self.assigned_by_var, _ = self.labeled_entry(tab, "Assigned By:")
        self.bin_no_var, _ = self.labeled_entry(tab, "Bin No:")
        self.researcher_var, _ = self.labeled_entry(tab, "Researcher Name:")
        self.total_hours_var, _ = self.labeled_entry(tab, "Total Hours:")
        symbar = ttk.Frame(tab); symbar.pack(fill="x", pady=(6, 0))
        ChemDropdown(symbar, self._target_widget).pack(side=tk.LEFT, padx=(0, 6))

    def _tab_technical(self, notebook):
        outer = ttk.Frame(notebook, padding=6)
        notebook.add(outer, text="2. Technical Information")
        self._tab_roots.append(outer); self._apply_ui_watermark(outer)
        sub = ttk.Notebook(outer); sub.pack(fill="both", expand=True)

        # 2.1.
        t21 = ttk.Frame(sub, padding=8); sub.add(t21, text="2.1. Plain Summary")
        self.txt_plain = self.labeled_text(
            t21, "2.1. Plain Language Summary:", height=10,
            ii_text="Plain-English overview of the customer's request and the problem to solve."
        )
        ChemDropdown(t21, self._target_widget).pack(anchor="w", pady=2)

        # 2.2. Objectives
        t22 = ttk.Frame(sub, padding=8); sub.add(t22, text="2.2. Objectives")
        info_row = ttk.Frame(t22); info_row.pack(fill="x")
        ttk.Label(info_row, text="2.2. Objectives (numbered):", width=30, anchor="w").pack(side=tk.LEFT)
        info_icon(info_row, "List specific, actionable objectives—one per line.")
        self._obj_container = ttk.Frame(t22); self._obj_container.pack(fill="both", expand=True, pady=(2, 0))
        self.obj_vars: List[tk.StringVar] = []; self._add_objective_row()
        btns = ttk.Frame(t22); btns.pack(anchor="w", pady=4)
        ttk.Button(btns, text="Add Objective", command=self._add_objective_row).pack(side=tk.LEFT)
        ttk.Button(btns, text="Remove Last", command=self._remove_last_objective).pack(side=tk.LEFT, padx=6)
        ChemDropdown(t22, self._target_widget).pack(anchor="w", pady=2)

        # 2.3. Methods
        t23 = ttk.Frame(sub, padding=8); sub.add(t23, text="2.3. Methods")

        frm_rm = ttk.Labelframe(t23, text="2.3.1. Raw Materials")
        frm_rm.pack(fill="both", expand=True, pady=4)
        self.txt_rm = self.labeled_text(frm_rm, "Items (one per line):", height=6,
                                        ii_text="List raw materials as bullet points.")
        ChemDropdown(frm_rm, self._target_widget).pack(anchor="w", pady=2)

        frm_ins = ttk.Labelframe(t23, text="2.3.2. Instrument")
        frm_ins.pack(fill="both", expand=True, pady=4)
        self.txt_ins = self.labeled_text(frm_ins, "Items (one per line):", height=6)

        frm_proc = ttk.Labelframe(t23, text="2.3.3. Experimental Procedure")
        frm_proc.pack(fill="both", expand=True, pady=4)
        self.txt_proc = self.labeled_text(
            frm_proc, "Steps (one per line):", height=8,
            ii_text="Write reproducible, step-by-step procedures."
        )
        ChemDropdown(frm_proc, self._target_widget).pack(anchor="w", pady=2)

        # 2.3.4. Trial History
        frm_trial = ttk.Labelframe(t23, text="2.3.4. Trial History")
        frm_trial.pack(fill="both", expand=True, pady=4)

        topbar = ttk.Frame(frm_trial); topbar.pack(fill="x", pady=(4, 6))
        ttk.Label(topbar, text="Layout:").pack(side=tk.LEFT)
        self.trial_layout_var = tk.StringVar(value=TRIAL_LAYOUTS[0])
        lay_cb = ttk.Combobox(topbar, textvariable=self.trial_layout_var, values=TRIAL_LAYOUTS,
                              state="readonly", width=16); lay_cb.pack(side=tk.LEFT, padx=4)
        ttk.Label(topbar, text="Word Style:").pack(side=tk.LEFT, padx=(8, 0))
        self.trial_style_var = tk.StringVar(value=DOCX_TABLE_STYLES[0])
        style_cb = ttk.Combobox(topbar, textvariable=self.trial_style_var, values=DOCX_TABLE_STYLES,
                                state="readonly", width=22); style_cb.pack(side=tk.LEFT, padx=4)
        ttk.Button(topbar, text="Delete Selected", command=self._trial_delete_selected).pack(side=tk.LEFT, padx=8)
        ttk.Button(topbar, text="Clear All", command=self._trial_clear_all).pack(side=tk.LEFT)
        info_icon(topbar, "Editable table; choose a layout and style. Column widths update with 'Layout'.")

        add_frame = ttk.Frame(frm_trial); add_frame.pack(fill="x", pady=(2, 6))
        self.trial_no = tk.StringVar(); self.trial_issue = tk.StringVar(); self.trial_reasons = tk.StringVar()
        ttk.Label(add_frame, text="Trial#:").grid(row=0, column=0, sticky="w")
        ent_no = ttk.Entry(add_frame, textvariable=self.trial_no, width=12); ent_no.grid(row=0, column=1, padx=4)
        ttk.Label(add_frame, text="Issue:").grid(row=0, column=2, sticky="w")
        ent_issue = ttk.Entry(add_frame, textvariable=self.trial_issue, width=40); ent_issue.grid(row=0, column=3, padx=4)
        ttk.Label(add_frame, text="Possible Reasons:").grid(row=0, column=4, sticky="w")
        ent_reasons = ttk.Entry(add_frame, textvariable=self.trial_reasons, width=40); ent_reasons.grid(row=0, column=5, padx=4)
        self._watch_widget(ent_no); self._watch_widget(ent_issue); self._watch_widget(ent_reasons)
        ttk.Button(add_frame, text="Add Row", command=self._add_trial_row).grid(row=0, column=6, padx=8, sticky="w")

        cols = ("Trial#", "Issue", "Possible Reasons")
        self.trial_tree = ttk.Treeview(frm_trial, columns=cols, show="headings", height=8)
        for c in cols:
            self.trial_tree.heading(c, text=c); self.trial_tree.column(c, width=180, stretch=True)
        self.trial_tree.pack(fill="both", expand=True, padx=0, pady=(0, 4))
        self.trial_tree.bind("<Double-1>", self._trial_begin_edit)

        lay_cb.bind("<<ComboboxSelected>>", lambda _e: self._update_trial_preview())
        style_cb.bind("<<ComboboxSelected>>", lambda _e: self._update_trial_preview())
        self._update_trial_preview()

        # 2.4. Results
        t24 = ttk.Frame(sub, padding=8)
        sub.add(t24, text="2.4. Results")
        self.results_container = ttk.Frame(t24)
        self.results_container.pack(fill="both", expand=True)
        self._results: List[ResultItem] = []
        self._rebuild_results_list()
        ttk.Button(t24, text="Add Result Item", command=self._add_result_dialog).pack(anchor="w", pady=6)

        # 2.5. Conclusion
        t25 = ttk.Frame(sub, padding=8); sub.add(t25, text="2.5. Conclusion")
        self.txt_conclusion = self.labeled_text(
            t25, "2.5. Bullet points (one per line):", height=6,
            ii_text="Concise, outcome-focused statements tied to objectives."
        )
        ChemDropdown(t25, self._target_widget).pack(anchor="w", pady=2)

        # 2.6. Miscellaneous
        t26_misc = ttk.Frame(sub, padding=8); sub.add(t26_misc, text="2.6. Miscellaneous")
        self.txt_misc = self.labeled_text(t26_misc, "2.6. Text:", height=6)
        ChemDropdown(t26_misc, self._target_widget).pack(anchor="w", pady=2)

        # 2.7. References
        t27 = ttk.Frame(sub, padding=8); sub.add(t27, text="2.7. References")
        self.txt_refs = self.labeled_text(t27, "2.7. References (one per line):", height=6,
                                          ii_text="Provide citations for all sources referenced.")
        ChemDropdown(t27, self._target_widget).pack(anchor="w", pady=2)

    def _tab_regulatory(self, notebook):
        tab = ttk.Frame(notebook, padding=10)
        notebook.add(tab, text="3. Regulatory")
        self._tab_roots.append(tab); self._apply_ui_watermark(tab)

        f1 = ttk.Labelframe(tab, text="3.1. Application regulations"); f1.pack(fill="both", expand=True, pady=4)
        self.txt_regs = self.labeled_text(f1, "Text:", height=6,
            ii_text="Summarize applicable EPA, FDA, USDA, NSF/ANSI, AWWA, and related requirements.")
        f2 = ttk.Labelframe(tab, text="3.2. Label Requirements"); f2.pack(fill="both", expand=True, pady=4)
        self.txt_label_req = self.labeled_text(f2, "Text:", height=6, ii_text="Required statements, directions, warnings, and use instructions.")
        f3 = ttk.Labelframe(tab, text="3.3. Certification Requirements"); f3.pack(fill="both", expand=True, pady=4)
        self.txt_cert_req = self.labeled_text(f3, "Text:", height=6, ii_text="e.g., NSF/ANSI 60, OMRI, Kosher, Halal, other marks.")

    def _tab_scaleup(self, notebook):
        tab = ttk.Frame(notebook, padding=10)
        notebook.add(tab, text="4. Scale Up")
        self._tab_roots.append(tab); self._apply_ui_watermark(tab)

        # Two-column grid like section 6
        grid = ttk.Frame(tab)
        grid.pack(fill="both", expand=True)
        grid.columnconfigure(0, weight=1)
        grid.columnconfigure(1, weight=1)

        # 4.1. Manufacturing Order (spans both columns)
        f41 = ttk.Labelframe(grid, text="4.1. Manufacturing Order")
        f41.grid(row=0, column=0, columnspan=2, sticky="nsew", padx=4, pady=4)
        info_row = ttk.Frame(f41); info_row.pack(fill="x")
        ttk.Label(info_row, text="Steps (one per line):", width=24, anchor="w").pack(side=tk.LEFT)
        info_icon(info_row,
                  "List each manufacturing step separately: formulation, order of addition, mixing speed/time/temperature/"
                  "atmosphere, filtration, in-process/QC checks, and any factor that may cause deviations from the lab batch.")
        self._mo_container = ttk.Frame(f41); self._mo_container.pack(fill="both", expand=True, pady=(2, 0))
        self.mo_vars: List[tk.StringVar] = []
        self._add_mo_row()
        btns_mo = ttk.Frame(f41); btns_mo.pack(anchor="w", pady=4)
        ttk.Button(btns_mo, text="Add Step", command=self._add_mo_row).pack(side=tk.LEFT)
        ttk.Button(btns_mo, text="Remove Last", command=self._remove_last_mo).pack(side=tk.LEFT, padx=6)
        ChemDropdown(f41, self._target_widget).pack(anchor="w", pady=2)

        # 4.2. Formulation Risk (left)
        f42 = ttk.Labelframe(grid, text="4.2. Formulation Risk")
        f42.grid(row=1, column=0, sticky="nsew", padx=4, pady=4)
        ii42 = ("Identify formulation risks that could cause product failure, regulatory noncompliance, or safety incidents. "
                "Note likelihood, impact, and proposed mitigations where known.")
        self.txt_42 = self.plain_text_area(f42, height=10, ii_text=ii42)

        # 4.3. Hazards (right)
        f43 = ttk.Labelframe(grid, text="4.3. Hazards")
        f43.grid(row=1, column=1, sticky="nsew", padx=4, pady=4)
        ii43 = ("Summarize process hazards: flash point, reactivity, corrosivity, fumes/vapors, dusting, exotherms, "
                "pressure build-up, and incompatibilities.")
        self.txt_43 = self.plain_text_area(f43, height=10, ii_text=ii43)

        # 4.4. Equipment (left)
        f44 = ttk.Labelframe(grid, text="4.4. Equipment")
        f44.grid(row=2, column=0, sticky="nsew", padx=4, pady=4)
        ii44 = ("List unique equipment needs: heating/cooling capacity, high-shear or inline mixing, inert gas blanketing, "
                "filtration/centrifugation, CIP/SIP, explosion-proofing, or other requirements.")
        self.txt_44 = self.plain_text_area(f44, height=10, ii_text=ii44)

        # 4.5. CAPEX Requirements (right)
        f45 = ttk.Labelframe(grid, text="4.5. CAPEX Requirements")
        f45.grid(row=2, column=1, sticky="nsew", padx=4, pady=4)
        ii45 = ("Estimate fixed assets to add or upgrade—buildings, utilities, vehicles, equipment, or land—and provide rough "
                "costs and timing if available.")
        self.txt_45 = self.plain_text_area(f45, height=10, ii_text=ii45)

        # 4.6. Safety Assessment (spans both)
        f46 = ttk.Labelframe(grid, text="4.6. Safety Assessment")
        f46.grid(row=3, column=0, columnspan=2, sticky="nsew", padx=4, pady=4)
        ii46 = ("Specify PPE, ventilation or containment, storage and segregation constraints, spill response, training, and any "
                "permits/approvals required.")
        self.txt_46 = self.plain_text_area(f46, height=8, ii_text=ii46)

    def _tab_quality(self, notebook):
        tab = ttk.Frame(notebook, padding=10)
        notebook.add(tab, text="5. Quality")
        self._tab_roots.append(tab); self._apply_ui_watermark(tab)

        # Two-column grid like section 6
        grid = ttk.Frame(tab)
        grid.pack(fill="both", expand=True)
        grid.columnconfigure(0, weight=1)
        grid.columnconfigure(1, weight=1)

        f1 = ttk.Labelframe(grid, text="5.1. Raw Material Sourcing")
        f1.grid(row=0, column=0, sticky="nsew", padx=4, pady=4)
        self.txt_q_sourcing = self.labeled_text(f1, "Text:", height=10,
                                                ii_text="Approved vendors, alternates, and supply risks.")

        f2 = ttk.Labelframe(grid, text="5.2. LIMS Setup")
        f2.grid(row=0, column=1, sticky="nsew", padx=4, pady=4)
        self.txt_q_lims = self.labeled_text(f2, "Text:", height=10,
                                            ii_text="Parameters, methods, specs, sampling plans, and release criteria.")

        f3 = ttk.Labelframe(grid, text="5.3. Stability Testing")
        f3.grid(row=1, column=0, sticky="nsew", padx=4, pady=4)
        self.txt_q_stability = self.labeled_text(f3, "Text:", height=10,
                                                 ii_text="Study design, storage conditions, timepoints, and acceptance ranges.")

        f4 = ttk.Labelframe(grid, text="5.4. Packaging Compatibility")
        f4.grid(row=1, column=1, sticky="nsew", padx=4, pady=4)
        self.txt_q_pack = self.labeled_text(f4, "Text:", height=10,
                                            ii_text="Container/closure, liners, valves, migration/leachables, and shelf life.")

        symbar = ttk.Frame(tab); symbar.pack(fill="x", pady=(6, 0))
        ChemDropdown(symbar, self._target_widget).pack(side=tk.LEFT, padx=(0, 6))

    def _small_text(self, parent, label, ii=None):
        frame = ttk.Frame(parent)
        top = ttk.Frame(frame); top.pack(fill="x")
        ttk.Label(top, text=label, anchor="w").pack(side=tk.LEFT)
        if ii: info_icon(top, ii)
        txt = tk.Text(frame, height=4, wrap="word")
        txt.pack(fill="both", expand=True)
        self._watch_widget(txt)
        return frame, txt

    def _smart_goal_row(self, parent, key, title, ii_text):
        row = ttk.Frame(parent)
        ttk.Label(row, text=f"{title}:", width=12, anchor="w").pack(side=tk.LEFT)
        info_icon(row, ii_text)
        var = tk.StringVar()
        ent = ttk.Entry(row, textvariable=var); ent.pack(side=tk.LEFT, fill="x", expand=True, padx=4)
        self._watch_widget(ent)
        return row, var

    def _tab_commercial(self, notebook):
        tab = ttk.Frame(notebook, padding=10)
        notebook.add(tab, text="6. Commercial")
        self._tab_roots.append(tab); self._apply_ui_watermark(tab)

        # Two-column grid container
        grid = ttk.Frame(tab); grid.pack(fill="both", expand=True)
        grid.columnconfigure(0, weight=1); grid.columnconfigure(1, weight=1)

        # 6.1. Customer Objectives / Problem Statement
        f61, self.txt_c_obj_problem = self._small_text(
            grid, "6.1. Customer Objectives / Problem Statement",
            ii="Plain-English summary of the customer's challenge, desired outcome, and target application."
        ); f61.grid(row=0, column=0, sticky="nsew", padx=4, pady=4)

        # 6.2. SMART Success Criteria
        f62 = ttk.Labelframe(grid, text="6.2. SMART Success Criteria"); f62.grid(row=0, column=1, sticky="nsew", padx=4, pady=4)
        self.smart_vars = {}
        rows = [
            ("S", "Specific", "State exactly what will be delivered or improved."),
            ("M", "Measurable", "Define how success will be quantified or verified."),
            ("A", "Achievable", "Explain why this is realistic given constraints."),
            ("R", "Relevant", "Show alignment with the customer's/business goals."),
            ("T", "Time-bound", "Provide a clear deadline or timeframe."),
        ]
        for i,(k,title,ii) in enumerate(rows):
            r, var = self._smart_goal_row(f62, k, title, ii); r.pack(fill="x", pady=2); self.smart_vars[k]=var

        # 6.3.–6.11. (revised ii content)
        f63, self.txt_c_specs = self._small_text(grid, "6.3. Customer Specifications",
                                                 ii="Performance specs, test methods, and acceptance criteria."); f63.grid(row=1, column=0, sticky="nsew", padx=4, pady=4)
        f64, self.txt_c_expected_volume = self._small_text(grid, "6.4. Expected Business Volume",
                                                           ii="Estimated annual/seasonal demand with units and key assumptions."); f64.grid(row=1, column=1, sticky="nsew", padx=4, pady=4)
        f65, self.txt_c_packaging_req = self._small_text(grid, "6.5. Packaging Requirement",
                                                         ii="Sizes, materials, closures, labeling, and any special handling/UN/DOT needs."); f65.grid(row=2, column=0, sticky="nsew", padx=4, pady=4)
        f66, self.txt_c_raw_material_prefs = self._small_text(grid, "6.6. Raw Material Restrictions / Preferences",
                                                              ii="Forbidden/allowed substances, required certifications (OMRI, Kosher, Halal), origin or purity constraints."); f66.grid(row=2, column=1, sticky="nsew", padx=4, pady=4)
        f67, self.txt_c_sample_needed = self._small_text(grid, "6.7. Sample Needed",
                                                         ii="Quantity, packaging, requested tests, and needed-by date."); f67.grid(row=3, column=0, sticky="nsew", padx=4, pady=4)
        f68, self.txt_c_opportunity_timeline = self._small_text(grid, "6.8. Opportunity Timeline",
                                                                ii="Key milestones: ITNO/NDA, sample ship, pilot/field trial, PO, launch."); f68.grid(row=3, column=1, sticky="nsew", padx=4, pady=4)
        f69, self.txt_c_target = self._small_text(grid, "6.9. Target Application",
                                                  ii="Crop/use case, application method, dose rates, and environmental conditions."); f69.grid(row=4, column=0, sticky="nsew", padx=4, pady=4)
        f610, self.txt_c_feedback = self._small_text(grid, "6.10. Customer Feedback",
                                                     ii="Notes from calls, demos, and trials—include quotes or quantified results when useful."); f610.grid(row=4, column=1, sticky="nsew", padx=4, pady=4)
        f611, self.txt_c_tds = self._small_text(grid, "6.11. TDS Development",
                                                ii="Status of the Technical Data Sheet—version, pending tests, owner, next steps."); f611.grid(row=5, column=0, sticky="nsew", padx=4, pady=4)

        # 6.12. Email Correspondence (table + New dialog)
        f612 = ttk.Labelframe(grid, text="6.12. Email Correspondence")
        f612.grid(row=5, column=1, sticky="nsew", padx=4, pady=4)
        cols = ("Date", "Customer Name", "Correspondence")
        self.email_tree = ttk.Treeview(f612, columns=cols, show="headings", height=6)
        for i, c in enumerate(cols):
            self.email_tree.heading(c, text=c)
            self.email_tree.column(c, width=(120 if i<2 else 420), stretch=True, anchor="w")
        self.email_tree.pack(fill="both", expand=True, pady=(0,4))
        btnrow = ttk.Frame(f612); btnrow.pack(fill="x")
        ttk.Button(btnrow, text="New", command=self._email_new_dialog).pack(side=tk.LEFT)
        ttk.Button(btnrow, text="Delete Selected", command=self._email_delete).pack(side=tk.LEFT, padx=6)

        symbar = ttk.Frame(tab); symbar.pack(fill="x", pady=(6, 0))
        ChemDropdown(symbar, self._target_widget).pack(side=tk.LEFT, padx=(0, 6))

    def _email_new_dialog(self):
        win = tk.Toplevel(self); win.title("New Email Entry"); win.geometry("640x420")
        frm = ttk.Frame(win, padding=8); frm.pack(fill="both", expand=True)

        ttk.Label(frm, text="Date (YYYY-MM-DD):").grid(row=0, column=0, sticky="w")
        v_date = tk.StringVar(value=datetime.date.today().isoformat()); ent_date = ttk.Entry(frm, textvariable=v_date); ent_date.grid(row=0, column=1, sticky="ew", padx=6, pady=4)

        ttk.Label(frm, text="Customer Name:").grid(row=1, column=0, sticky="w")
        v_cust = tk.StringVar(); ent_cust = ttk.Entry(frm, textvariable=v_cust); ent_cust.grid(row=1, column=1, sticky="ew", padx=6, pady=4)

        ttk.Label(frm, text="Correspondence:").grid(row=2, column=0, sticky="nw")
        txt_corr = tk.Text(frm, height=10, wrap="word"); txt_corr.grid(row=2, column=1, sticky="nsew", padx=6, pady=4)

        frm.columnconfigure(1, weight=1); frm.rowconfigure(2, weight=1)

        def save():
            d = v_date.get().strip(); c = v_cust.get().strip(); t = txt_corr.get("1.0","end").strip()
            if not (d and c and t):
                messagebox.showwarning("Missing", "Please enter Date, Customer Name, and Correspondence.")
                return
            self.email_tree.insert("", "end", values=(d, c, t))
            win.destroy()
        btns = ttk.Frame(frm); btns.grid(row=3, column=1, sticky="e", pady=6)
        ttk.Button(btns, text="Save", command=save).pack(side=tk.LEFT, padx=4)
        ttk.Button(btns, text="Cancel", command=win.destroy).pack(side=tk.LEFT, padx=4)

    def _email_delete(self):
        for iid in self.email_tree.selection():
            self.email_tree.delete(iid)

    def _tab_export(self, notebook):
        tab = ttk.Frame(notebook, padding=10)
        notebook.add(tab, text="Export / Generate")
        self._tab_roots.append(tab); self._apply_ui_watermark(tab)

        hdr = ttk.Frame(tab); hdr.pack(fill="x", pady=6)
        ttk.Label(hdr, text="Header image (PNG/JPG):").pack(side=tk.LEFT)
        ttk.Entry(hdr, textvariable=self.header_path_var, width=80).pack(side=tk.LEFT, padx=6)
        ttk.Button(hdr, text="Browse...", command=self._pick_header).pack(side=tk.LEFT)

        ftr = ttk.Frame(tab); ftr.pack(fill="x", pady=6)
        ttk.Label(ftr, text="Footer image (PNG/JPG):").pack(side=tk.LEFT)
        ttk.Entry(ftr, textvariable=self.footer_path_var, width=80).pack(side=tk.LEFT, padx=6)
        ttk.Button(ftr, text="Browse...", command=self._pick_footer).pack(side=tk.LEFT)

        fmt = ttk.Frame(tab); fmt.pack(fill="x", pady=6)
        ttk.Label(fmt, text="Export format:").pack(side=tk.LEFT)
        self.export_fmt = tk.StringVar(value="both")
        for v, lab in [("docx", "DOCX"), ("pdf", "PDF"), ("both", "Both")]:
            ttk.Radiobutton(fmt, text=lab, variable=self.export_fmt, value=v).pack(side=tk.LEFT, padx=6)

        chooser = ttk.Labelframe(tab, text="Include in output"); chooser.pack(fill="both", expand=True, pady=(6, 10))
        self.include = {}
        def mkvar(key, default=True):
            self.include[key] = tk.BooleanVar(value=default); return self.include[key]

        sec_frame = ttk.Frame(chooser); sec_frame.pack(fill="x", pady=(6, 2))
        ttk.Label(sec_frame, text="1. General").grid(row=0, column=0, sticky="w")
        ttk.Checkbutton(sec_frame, text="Include", variable=mkvar("sec_general")).grid(row=0, column=1, sticky="w", padx=8)

        tech = ttk.LabelFrame(chooser, text="2. Technical")
        tech.pack(fill="x", padx=4, pady=4)
        ttk.Checkbutton(tech, text="2.1. Plain Summary", variable=mkvar("t_plain")).pack(anchor="w")
        ttk.Checkbutton(tech, text="2.2. Objectives", variable=mkvar("t_objectives")).pack(anchor="w")
        ttk.Checkbutton(tech, text="2.3. Methods", variable=mkvar("t_methods")).pack(anchor="w")
        ttk.Checkbutton(tech, text="  2.3.1. Raw Materials", variable=mkvar("t_rm")).pack(anchor="w", padx=16)
        ttk.Checkbutton(tech, text="  2.3.2. Instrument", variable=mkvar("t_ins")).pack(anchor="w", padx=16)
        ttk.Checkbutton(tech, text="  2.3.3. Experimental Procedure", variable=mkvar("t_proc")).pack(anchor="w", padx=16)
        ttk.Checkbutton(tech, text="  2.3.4. Trial History", variable=mkvar("t_trial")).pack(anchor="w", padx=16)
        ttk.Checkbutton(tech, text="2.4. Results", variable=mkvar("t_results")).pack(anchor="w")
        ttk.Checkbutton(tech, text="2.5. Conclusion", variable=mkvar("t_conc")).pack(anchor="w")
        ttk.Checkbutton(tech, text="2.6. Miscellaneous", variable=mkvar("t_misc")).pack(anchor="w")
        ttk.Checkbutton(tech, text="2.7. References", variable=mkvar("t_refs")).pack(anchor="w")

        ttk.Checkbutton(chooser, text="3. Regulatory", variable=mkvar("sec_reg")).pack(anchor="w")
        ttk.Checkbutton(chooser, text="4. Scale Up", variable=mkvar("sec_scale")).pack(anchor="w")
        ttk.Checkbutton(chooser, text="5. Quality", variable=mkvar("sec_quality")).pack(anchor="w")
        ttk.Checkbutton(chooser, text="6. Commercial", variable=mkvar("sec_commercial")).pack(anchor="w")

        out = ttk.Frame(tab); out.pack(fill="x", pady=6)
        ttk.Label(out, text="Output folder:").pack(side=tk.LEFT)
        self.out_dir_var = tk.StringVar(value=os.path.join(app_writable_dir(), "reports"))
        ttk.Entry(out, textvariable=self.out_dir_var, width=60).pack(side=tk.LEFT, padx=6)
        ttk.Button(out, text="Choose...", command=self._pick_outdir).pack(side=tk.LEFT)

        ttk.Button(tab, text="Generate Report", command=self._generate).pack(anchor="w", pady=12)
        ttk.Label(tab, text="Tip: For 2.4. tables, paste directly from Excel/CSV—I'll detect the columns.",
                  foreground="#555").pack(anchor="w", pady=6)

        for k in ["sec_general", "t_plain", "t_objectives", "t_methods", "t_rm", "t_ins", "t_proc", "t_trial",
                  "t_results", "t_conc", "t_misc", "t_refs", "sec_reg", "sec_scale", "sec_quality", "sec_commercial"]:
            self.include[k].set(True)

    # ---------- Objectives rows ----------
    def _add_objective_row(self):
        idx = len(self.obj_vars) + 1
        row = ttk.Frame(self._obj_container); row.pack(fill="x", pady=2)
        ttk.Label(row, text=f"{idx}.", width=3).pack(side=tk.LEFT)
        var = tk.StringVar()
        ent = ttk.Entry(row, textvariable=var); ent.pack(side=tk.LEFT, fill="x", expand=True)
        self._watch_widget(ent)
        self.obj_vars.append(var)
    def _remove_last_objective(self):
        if not self._obj_container.winfo_children(): return
        self._obj_container.winfo_children()[-1].destroy()
        if self.obj_vars: self.obj_vars.pop()

    # ---------- Manufacturing Order rows (4.1.) ----------
    def _add_mo_row(self):
        idx = len(getattr(self, "mo_vars", [])) + 1
        row = ttk.Frame(self._mo_container); row.pack(fill="x", pady=2)
        ttk.Label(row, text=f"{idx}.", width=3).pack(side=tk.LEFT)
        var = tk.StringVar()
        ent = ttk.Entry(row, textvariable=var)
        ent.pack(side=tk.LEFT, fill="x", expand=True)
        self._watch_widget(ent)
        self.mo_vars.append(var)
    def _remove_last_mo(self):
        if not self._mo_container.winfo_children(): return
        self._mo_container.winfo_children()[-1].destroy()
        if self.mo_vars: self.mo_vars.pop()

    # ---------- Trial history helpers ----------
    def _add_trial_row(self):
        num = self.trial_no.get().strip(); issue = self.trial_issue.get().strip(); reasons = self.trial_reasons.get().strip()
        if not num:
            messagebox.showwarning("Missing", "Please enter Trial#."); return
        self.trial_tree.insert("", "end", values=(num, issue, reasons))
        self.trial_no.set(""); self.trial_issue.set(""); self.trial_reasons.set("")
    def _trial_delete_selected(self):
        for iid in self.trial_tree.selection():
            self.trial_tree.delete(iid)
    def _trial_clear_all(self):
        for iid in self.trial_tree.get_children(""):
            self.trial_tree.delete(iid)
    def _trial_begin_edit(self, event):
        if self.trial_tree.identify("region", event.x, event.y) != "cell": return
        row_id = self.trial_tree.identify_row(event.y); col_id = self.trial_tree.identify_column(event.x)
        if not (row_id and col_id): return
        x,y,w,h = self.trial_tree.bbox(row_id, col_id)
        col_index = int(col_id.replace("#","")) - 1
        old = self.trial_tree.item(row_id, "values")[col_index]
        edit = tk.Entry(self.trial_tree); edit.insert(0, old); edit.place(x=x, y=y, width=w, height=h)
        def save_edit(_e=None):
            vals = list(self.trial_tree.item(row_id, "values")); vals[col_index] = edit.get()
            self.trial_tree.item(row_id, values=vals); edit.destroy()
        edit.bind("<Return>", save_edit); edit.bind("<FocusOut>", lambda _e: save_edit()); edit.focus_set()
    def _update_trial_preview(self):
        tree = getattr(self, "trial_tree", None)
        if not tree: return
        lay = self.trial_layout_var.get()
        if lay == "Even columns": widths = (120, 280, 280)
        elif lay == "Reasons wide": widths = (100, 200, 400)
        else: widths = (90, 220, 310)
        for w, col in zip(widths, ("#1","#2","#3")):
            tree.column(col, width=int(w * 0.9), stretch=True)

    # ---------- Results list & dialogs ----------
    def _rebuild_results_list(self):
        for ch in self.results_container.winfo_children():
            ch.destroy()
        if not getattr(self, "_results", None):
            ttk.Label(self.results_container, text="No results added yet.").pack(anchor="w")
            return
        for i, it in enumerate(self._results, 1):
            row = ttk.Frame(self.results_container); row.pack(fill="x", pady=3)
            kind_label = it.kind.upper(); extra = ""
            if it.kind == "image":
                count = len(it.images) if it.images else (1 if it.image_path else 0)
                extra = f" — {count} image(s)"
            ttk.Label(row, text=f"{i}. {it.title}  [{kind_label}]{extra}").pack(side=tk.LEFT)
            ttk.Button(row, text="Edit", command=lambda idx=i-1: self._edit_result_dialog(idx)).pack(side=tk.LEFT, padx=6)
            ttk.Button(row, text="Remove", command=lambda idx=i-1: self._remove_result(idx)).pack(side=tk.LEFT)

    def _remove_result(self, idx: int):
        if 0 <= idx < len(self._results):
            del self._results[idx]; self._rebuild_results_list()

    def _add_result_dialog(self): 
        self._result_dialog()

    def _edit_result_dialog(self, idx: int):
        if 0 <= idx < len(self._results): 
            self._result_dialog(self._results[idx], idx)

    def _result_dialog(self, existing: Optional[ResultItem]=None, edit_index: Optional[int]=None):
        win = tk.Toplevel(self); win.title("Result Item"); win.geometry("820x560")
        frm = ttk.Frame(win, padding=8); frm.pack(fill="both", expand=True)

        # Title
        ttk.Label(frm, text="Title:").grid(row=0, column=0, sticky="w")
        v_title = tk.StringVar(value=(existing.title if existing else ""))
        ent_title = ttk.Entry(frm, textvariable=v_title); ent_title.grid(row=0, column=1, sticky="ew", padx=6, pady=4)

        # Kind
        ttk.Label(frm, text="Type:").grid(row=1, column=0, sticky="w")
        v_kind = tk.StringVar(value=(existing.kind if existing else "text"))
        kinds = ttk.Frame(frm); kinds.grid(row=1, column=1, sticky="w", padx=6, pady=4)
        for val, lab in [("text","Text"),("table","Table"),("image","Image(s)")]:
            ttk.Radiobutton(kinds, text=lab, value=val, variable=v_kind).pack(side=tk.LEFT, padx=(0,8))

        # Stacked frames
        stack = ttk.Frame(frm); stack.grid(row=2, column=0, columnspan=2, sticky="nsew")
        frm.rowconfigure(2, weight=1); frm.columnconfigure(1, weight=1)

        # TEXT frame
        text_frame = ttk.Frame(stack); 
        txt_text = tk.Text(text_frame, height=16, wrap="word"); txt_text.pack(fill="both", expand=True)
        if existing and existing.kind=="text":
            txt_text.insert("1.0", existing.content)

        # TABLE frame
        table_frame = ttk.Frame(stack)
        ttk.Label(table_frame, text="Paste CSV/TSV or plain text table (auto-detect):").pack(anchor="w")
        txt_table = tk.Text(table_frame, height=14, wrap="none"); txt_table.pack(fill="both", expand=True, pady=(2,4))
        style_row = ttk.Frame(table_frame); style_row.pack(fill="x")
        ttk.Label(style_row, text="Word table style:").pack(side=tk.LEFT)
        v_tbl_style = tk.StringVar(value=(existing.table_style if existing else "Table Grid"))
        ttk.Combobox(style_row, textvariable=v_tbl_style, values=DOCX_TABLE_STYLES, state="readonly", width=26).pack(side=tk.LEFT, padx=6)
        if existing and existing.kind=="table":
            if existing.table_data:
                txt_table.insert("1.0", "\n".join("\t".join(r) for r in existing.table_data))
            else:
                txt_table.insert("1.0", existing.content or "")

        # IMAGE frame
        image_frame = ttk.Frame(stack)
        mgr = ImageManager(image_frame, initial_files=(existing.images if existing and existing.kind=="image" else None))
        mgr.pack(fill="both", expand=True)
        cap_row = ttk.Frame(image_frame); cap_row.pack(fill="x", pady=4)
        ttk.Label(cap_row, text="Caption:").pack(side=tk.LEFT)
        v_caption = tk.StringVar(value=(existing.caption if existing and existing.kind=="image" else ""))
        ttk.Entry(cap_row, textvariable=v_caption).pack(side=tk.LEFT, fill="x", expand=True, padx=6)

        frames = {"text": text_frame, "table": table_frame, "image": image_frame}
        def show(which):
            for k,f in frames.items():
                f.forget()
            frames[which].pack(fill="both", expand=True)
        show(v_kind.get())
        v_kind.trace_add("write", lambda *_: show(v_kind.get()))

        # Buttons
        btns = ttk.Frame(frm); btns.grid(row=3, column=1, sticky="e", pady=8)
        def save():
            title = v_title.get().strip() or "Untitled"
            kind = v_kind.get()
            if kind == "text":
                content = txt_text.get("1.0","end").strip()
                item = ResultItem(title=title, kind="text", content=content)
            elif kind == "table":
                raw = txt_table.get("1.0","end").strip()
                data = parse_table_text(raw)
                item = ResultItem(title=title, kind="table", content=raw, table_data=data, table_style=v_tbl_style.get())
            else:  # image
                imgs = mgr.save_edited_to_temp() if mgr.get_files() else []
                item = ResultItem(title=title, kind="image", images=imgs, caption=v_caption.get().strip())
            if edit_index is None:
                self._results.append(item)
            else:
                self._results[edit_index] = item
            self._rebuild_results_list()
            win.destroy()
        ttk.Button(btns, text="Save", command=save).pack(side=tk.LEFT, padx=4)
        ttk.Button(btns, text="Cancel", command=win.destroy).pack(side=tk.LEFT, padx=4)

    # ---------- Export helpers ----------
    def _pick_header(self):
        p = filedialog.askopenfilename(title="Select header image", filetypes=[("Images", "*.png;*.jpg;*.jpeg")])
        if p: self.header_path_var.set(p)
    def _pick_footer(self):
        p = filedialog.askopenfilename(title="Select footer image", filetypes=[("Images", "*.png;*.jpg;*.jpeg")])
        if p: self.footer_path_var.set(p)
    def _pick_outdir(self):
        p = filedialog.askdirectory(title="Select output folder")
        if p: self.out_dir_var.set(p)

    def _collect(self) -> ReportModel:
        r = ReportModel()
        r.project_title = self.ent_title.get().strip()
        r.start_date = self.start_date_var.get().strip()
        r.report_date = self.report_date_var.get().strip()
        r.assigned_by = self.assigned_by_var.get().strip()
        r.bin_no = self.bin_no_var.get().strip()
        r.researcher_name = self.researcher_var.get().strip()
        r.total_hours = self.total_hours_var.get().strip()

        r.plain_summary = self.txt_plain.get("1.0", "end").strip()
        r.objectives = [v.get().strip() for v in self.obj_vars if v.get().strip()]
        r.methods_raw_materials = safe_split_lines(self.txt_rm.get("1.0", "end"))
        r.methods_instruments = safe_split_lines(self.txt_ins.get("1.0", "end"))
        r.methods_procedure = safe_split_lines(self.txt_proc.get("1.0", "end"))

        r.trial_history = []
        for child in self.trial_tree.get_children(""):
            vals = self.trial_tree.item(child)["values"]
            r.trial_history.append(TrialRow(number=str(vals[0]), issue=str(vals[1]), reasons=str(vals[2])))
        r.trial_layout = self.trial_layout_var.get()
        r.trial_docx_style = self.trial_style_var.get()

        r.results = list(self._results)
        r.conclusion = safe_split_lines(self.txt_conclusion.get("1.0", "end"))
        r.miscellaneous = self.txt_misc.get("1.0", "end").strip()
        r.references = safe_split_lines(self.txt_refs.get("1.0", "end"))

        # Regulatory
        r.regulations = self.txt_regs.get("1.0", "end").strip()
        r.label_req = self.txt_label_req.get("1.0", "end").strip()
        r.certification_req = self.txt_cert_req.get("1.0", "end").strip()

        # Scale Up (4.1.–4.6.)
        r.manuf_order_steps = [v.get().strip() for v in self.mo_vars if v.get().strip()]
        r.formulation_risk_text = self.txt_42.get("1.0","end").strip()
        r.hazards_text = self.txt_43.get("1.0","end").strip()
        r.equipment_text = self.txt_44.get("1.0","end").strip()
        r.capex_text = self.txt_45.get("1.0","end").strip()
        r.safety_assess_text = self.txt_46.get("1.0","end").strip()

        # Quality
        r.raw_material_sourcing = self.txt_q_sourcing.get("1.0", "end").strip()
        r.lims_setup = self.txt_q_lims.get("1.0", "end").strip()
        r.stability_testing = self.txt_q_stability.get("1.0", "end").strip()
        r.packaging_compatibility = self.txt_q_pack.get("1.0", "end").strip()

        # Commercial (6.x.)
        r.c_obj_problem = self.txt_c_obj_problem.get("1.0","end").strip()
        r.smart_goals = {k:self.smart_vars[k].get().strip() for k in self.smart_vars}
        r.c_specs = self.txt_c_specs.get("1.0","end").strip()
        r.c_expected_volume = self.txt_c_expected_volume.get("1.0","end").strip()
        r.c_packaging_req = self.txt_c_packaging_req.get("1.0","end").strip()
        r.c_raw_material_prefs = self.txt_c_raw_material_prefs.get("1.0","end").strip()
        r.c_sample_needed = self.txt_c_sample_needed.get("1.0","end").strip()
        r.c_opportunity_timeline = self.txt_c_opportunity_timeline.get("1.0","end").strip()
        r.target_application = self.txt_c_target.get("1.0","end").strip()
        r.customer_feedback = self.txt_c_feedback.get("1.0","end").strip()
        r.tds_development = self.txt_c_tds.get("1.0","end").strip()

        r.email_correspondence = []
        for iid in self.email_tree.get_children(""):
            d,c,t = self.email_tree.item(iid, "values")
            r.email_correspondence.append(EmailEntry(date=str(d), customer=str(c), correspondence=str(t)))

        return r

    # ------------------------- DOCX helpers & export -------------------------
    def _H1(self, doc, text):
        p = doc.add_paragraph()
        run = p.add_run(text); run.bold = True; run.font.size = Pt(12)
        p.paragraph_format.left_indent = Inches(0)
        return p
    def _H2(self, doc, text):
        p = doc.add_paragraph()
        run = p.add_run(text); run.bold = True; run.font.size = Pt(11)
        p.paragraph_format.left_indent = Inches(0.25)  # subsection
        return p
    def _H3(self, doc, text):
        p = doc.add_paragraph()
        run = p.add_run(text); run.bold = True; run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.5)  # sub-subsection
        return p
    def _body_para(self, doc, text, level=2):
        p = doc.add_paragraph(text or "")
        p.paragraph_format.left_indent = Inches(0.25 if level==2 else 0.5)
        return p
    def _list_para(self, doc, text, numbered=False, level=2):
        p = doc.add_paragraph(text or "")
        p.style = doc.styles["List Number"] if numbered else doc.styles["List Bullet"]
        p.paragraph_format.left_indent = Inches(0.25 if level==2 else 0.5)
        return p

    def _add_page_number_footer_docx(self, doc, footer_img: Optional[str]):
        for section in doc.sections:
            # Footer image (bottom band)
            footer = section.footer
            # Clear default empty paragraph text
            if footer.paragraphs:
                footer.paragraphs[0].clear()
            if footer_img and os.path.isfile(footer_img):
                p_img = footer.add_paragraph()
                run_img = p_img.add_run()
                # Scale image to printable width
                try:
                    usable_w_in = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
                except Exception:
                    usable_w_in = 7.0
                try:
                    run_img.add_picture(footer_img, width=Inches(usable_w_in))
                except Exception:
                    pass
            # Page number at right
            p = footer.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE'); p._p.append(fld)

            # Place footer closer to the bottom edge
            try:
                section.footer_distance = Inches(0.25)
            except Exception:
                pass

    def _add_header_docx(self, doc, header_img: Optional[str]):
        for section in doc.sections:
            header = section.header
            # Clear default empty paragraph text
            if header.paragraphs:
                header.paragraphs[0].clear()
            if header_img and os.path.isfile(header_img):
                p_img = header.add_paragraph()
                run_img = p_img.add_run()
                try:
                    usable_w_in = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
                except Exception:
                    usable_w_in = 7.0
                try:
                    run_img.add_picture(header_img, width=Inches(usable_w_in))
                except Exception:
                    pass
            # Place header closer to the top edge
            try:
                section.header_distance = Inches(0.25)
            except Exception:
                pass

            # Ensure same header/footer on first page too
            try:
                section.different_first_page_header_footer = False
            except Exception:
                pass

    def _export_docx(self, r: ReportModel, out_path: str, header_img: Optional[str], footer_img: Optional[str]):
        include = self.include
        doc = Document()

        # Apply proper header & footer bands on all pages
        self._add_header_docx(doc, header_img)
        self._add_page_number_footer_docx(doc, footer_img)

        # Title (centered)
        p = doc.add_paragraph()
        run = p.add_run(r.project_title); run.bold = True; run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if include["sec_general"].get():
            self._H1(doc, "\n1. General Information")
            table = doc.add_table(rows=0, cols=2)
            def add_row(k, v):
                row = table.add_row().cells; row[0].text = k; row[1].text = v
            add_row("Start Date", r.start_date)
            add_row("Report Date", r.report_date)
            add_row("Assigned By", r.assigned_by)
            add_row("Bin No", r.bin_no)
            add_row("Researcher Name", r.researcher_name)
            add_row("Total Hours", r.total_hours)

        # 2. Technical
        self._H1(doc, "\n2. Technical Information")

        if include["t_plain"].get():
            self._H2(doc, "2.1. Plain Language Summary")
            self._body_para(doc, r.plain_summary, level=2)

        if include["t_objectives"].get():
            self._H2(doc, "2.2. Objectives")
            for it in r.objectives:
                self._list_para(doc, it, numbered=True, level=2)

        if include["t_methods"].get():
            self._H2(doc, "2.3. Methods")
            if include["t_rm"].get():
                self._H3(doc, "2.3.1. Raw Materials")
                for x in r.methods_raw_materials:
                    self._list_para(doc, x, numbered=False, level=3)
            if include["t_ins"].get():
                self._H3(doc, "2.3.2. Instrument")
                for x in r.methods_instruments:
                    self._list_para(doc, x, numbered=False, level=3)
            if include["t_proc"].get():
                self._H3(doc, "2.3.3. Experimental Procedure")
                for x in r.methods_procedure:
                    self._list_para(doc, x, numbered=False, level=3)
            if include["t_trial"].get():
                self._H3(doc, "2.3.4. Trial History")
                if r.trial_history:
                    t = doc.add_table(rows=1, cols=3)
                    try: t.style = r.trial_docx_style
                    except Exception: t.style = "Table Grid"
                    hdr = t.rows[0].cells
                    hdr[0].text, hdr[1].text, hdr[2].text = "Trial#", "Issue", "Possible Reasons"
                    for tr in r.trial_history:
                        c = t.add_row().cells
                        c[0].text, c[1].text, c[2].text = tr.number, tr.issue, tr.reasons

        if include["t_results"].get():
            self._H2(doc, "2.4. Results")
            for idx, it in enumerate(r.results, 1):
                if it.kind == "text":
                    self._H3(doc, f"2.4.{idx}. {it.title}")
                    for line in safe_split_lines(it.content):
                        self._body_para(doc, line, level=3)
                elif it.kind == "table":
                    self._H3(doc, f"{it.title}")
                    data = it.table_data if it.table_data else parse_table_text(it.content)
                    if data:
                        rows = len(data); cols = max(len(rw) for rw in data)
                        tb = doc.add_table(rows=rows, cols=cols)
                        try: tb.style = it.table_style
                        except Exception: tb.style = "Table Grid"
                        for r_i, row in enumerate(data):
                            for c_i in range(cols):
                                tb.cell(r_i, c_i).text = row[c_i] if c_i < len(row) else ""
                elif it.kind == "image":
                    imgs = it.images if it.images else ([it.image_path] if it.image_path else [])
                    if imgs:
                        if it.title: self._H3(doc, f"{it.title}")
                        for pth in imgs:
                            if pth and os.path.isfile(pth):
                                try: doc.add_picture(pth, width=Inches(5.8))
                                except Exception: continue
                                if it.caption:
                                    pcap = doc.add_paragraph(it.caption); 
                                    pcap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    pcap.paragraph_format.left_indent = Inches(0.5)

        if include["t_conc"].get():
            self._H2(doc, "2.5. Conclusion")
            for x in r.conclusion:
                self._list_para(doc, x, numbered=False, level=2)

        if include["t_misc"].get():
            self._H2(doc, "2.6. Miscellaneous")
            self._body_para(doc, r.miscellaneous, level=2)

        if include["t_refs"].get():
            self._H2(doc, "2.7. References")
            for x in r.references:
                self._list_para(doc, x, numbered=False, level=2)

        # 3. Regulatory
        if include["sec_reg"].get():
            self._H1(doc, "\n3. Regulatory")
            self._H2(doc, "3.1. Application regulations"); self._body_para(doc, r.regulations, level=2)
            self._H2(doc, "3.2. Label Requirements"); self._body_para(doc, r.label_req, level=2)
            self._H2(doc, "3.3. Certification Requirements"); self._body_para(doc, r.certification_req, level=2)

        # 4. Scale Up
        if include["sec_scale"].get():
            self._H1(doc, "\n4. Scale Up")
            self._H2(doc, "4.1. Manufacturing Order")
            for step in r.manuf_order_steps:
                self._list_para(doc, step, numbered=False, level=2)

            self._H2(doc, "4.2. Formulation Risk")
            self._body_para(doc, r.formulation_risk_text, level=2)

            self._H2(doc, "4.3. Hazards")
            self._body_para(doc, r.hazards_text, level=2)

            self._H2(doc, "4.4. Equipment")
            self._body_para(doc, r.equipment_text, level=2)

            self._H2(doc, "4.5. CAPEX Requirements")
            self._body_para(doc, r.capex_text, level=2)

            self._H2(doc, "4.6. Safety Assessment")
            self._body_para(doc, r.safety_assess_text, level=2)

        # 5. Quality
        if include["sec_quality"].get():
            self._H1(doc, "\n5. Quality")
            self._H2(doc, "5.1. Raw Material Sourcing"); self._body_para(doc, r.raw_material_sourcing, level=2)
            self._H2(doc, "5.2. LIMS Setup"); self._body_para(doc, r.lims_setup, level=2)
            self._H2(doc, "5.3. Stability Testing"); self._body_para(doc, r.stability_testing, level=2)
            self._H2(doc, "5.4. Packaging Compatibility"); self._body_para(doc, r.packaging_compatibility, level=2)

        # 6. Commercial
        if include["sec_commercial"].get():
            self._H1(doc, "\n6. Commercial")
            self._H2(doc, "6.1. Customer Objectives / Problem Statement"); self._body_para(doc, r.c_obj_problem, level=2)
            self._H2(doc, "6.2. SMART Success Criteria")
            for key, label in [("S","Specific"),("M","Measurable"),("A","Achievable"),("R","Relevant"),("T","Time-bound")]:
                if r.smart_goals.get(key,"").strip():
                    self._body_para(doc, f"{label}: {r.smart_goals[key]}", level=3)
            self._H2(doc, "6.3. Customer Specifications"); self._body_para(doc, r.c_specs, level=2)
            self._H2(doc, "6.4. Expected Business Volume"); self._body_para(doc, r.c_expected_volume, level=2)
            self._H2(doc, "6.5. Packaging Requirement"); self._body_para(doc, r.c_packaging_req, level=2)
            self._H2(doc, "6.6. Raw Material Restrictions / Preferences"); self._body_para(doc, r.c_raw_material_prefs, level=2)
            self._H2(doc, "6.7. Sample Needed"); self._body_para(doc, r.c_sample_needed, level=2)
            self._H2(doc, "6.8. Opportunity Timeline"); self._body_para(doc, r.c_opportunity_timeline, level=2)
            self._H2(doc, "6.9. Target Application"); self._body_para(doc, r.target_application, level=2)
            self._H2(doc, "6.10. Customer Feedback"); self._body_para(doc, r.customer_feedback, level=2)
            self._H2(doc, "6.11. TDS Development"); self._body_para(doc, r.tds_development, level=2)
            self._H2(doc, "6.12. Email Correspondence")
            if r.email_correspondence:
                t = doc.add_table(rows=1, cols=3)
                try: t.style = "Table Grid"
                except Exception: pass
                hdr = t.rows[0].cells
                hdr[0].text, hdr[1].text, hdr[2].text = "Date", "Customer Name", "Correspondence"
                for e in r.email_correspondence:
                    c = t.add_row().cells
                    c[0].text, c[1].text, c[2].text = e.date, e.customer, e.correspondence

        doc.save(out_path)

    # -------------------------- PDF export (ReportLab) ------------------------
    def _export_pdf_basic(self, r: ReportModel, out_path: str, header_img: Optional[str], footer_img: Optional[str]):
        if SimpleDocTemplate is None:
            messagebox.showerror("reportlab not found", "Install reportlab or enable docx2pdf for PDF export.")
            return

        include = self.include
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name="Title14", parent=styles["Title"], fontSize=14, leading=16,
            spaceBefore=0, spaceAfter=6, fontName="Helvetica-Bold", alignment=1
        ))
        styles.add(ParagraphStyle(name="H1_12", fontSize=12, leading=14, spaceBefore=6, spaceAfter=4, fontName="Helvetica-Bold", leftIndent=0))
        styles.add(ParagraphStyle(name="H2_11", fontSize=11, leading=13, spaceBefore=6, spaceAfter=4, fontName="Helvetica-Bold", leftIndent=18))  # 0.25"
        styles.add(ParagraphStyle(name="H3_10", fontSize=10, leading=12, spaceBefore=4, spaceAfter=2, fontName="Helvetica-Bold", leftIndent=36))  # 0.5"
        styles.add(ParagraphStyle(name="Body10_H2", fontSize=10, leading=12, spaceBefore=0, spaceAfter=4, leftIndent=18)) # align with H2
        styles.add(ParagraphStyle(name="Body10_H3", fontSize=10, leading=12, spaceBefore=0, spaceAfter=4, leftIndent=36)) # align with H3

        story = []

        # Title
        story.append(Paragraph(f"{r.project_title}", styles["Title14"]))
        story.append(Spacer(1, 0.12 * inch))

        if include["sec_general"].get():
            story.append(Paragraph("1. General Information", styles["H1_12"]))
            data = [
                ["Start Date", r.start_date],
                ["Report Date", r.report_date],
                ["Assigned By", r.assigned_by],
                ["Bin No", r.bin_no],
                ["Researcher Name", r.researcher_name],
                ["Total Hours", r.total_hours],
            ]
            t = Table(data, colWidths=[2.2 * inch, 4.8 * inch])
            t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.25, colors.grey)]))
            story.append(t); story.append(Spacer(1, 0.12 * inch))

        story.append(Paragraph("2. Technical Information", styles["H1_12"]))

        if include["t_plain"].get():
            story.append(Paragraph("2.1. Plain Language Summary", styles["H2_11"]))
            story.append(Paragraph((r.plain_summary or "").replace("\n", "<br/>"), styles["Body10_H2"]))

        if include["t_objectives"].get():
            story.append(Paragraph("2.2. Objectives", styles["H2_11"]))
            if r.objectives:
                story.append(ListFlowable([ListItem(Paragraph(o, styles["Body10_H2"])) for o in r.objectives], bulletType="1"))

        if include["t_methods"].get():
            story.append(Paragraph("2.3. Methods", styles["H2_11"]))
            if include["t_rm"].get():
                story.append(Paragraph("2.3.1. Raw Materials", styles["H3_10"]))
                story.append(ListFlowable([Paragraph(x, styles["Body10_H3"]) for x in r.methods_raw_materials], bulletType="bullet"))
            if include["t_ins"].get():
                story.append(Paragraph("2.3.2. Instrument", styles["H3_10"]))
                story.append(ListFlowable([Paragraph(x, styles["Body10_H3"]) for x in r.methods_instruments], bulletType="bullet"))
            if include["t_proc"].get():
                story.append(Paragraph("2.3.3. Experimental Procedure", styles["H3_10"]))
                story.append(ListFlowable([Paragraph(x, styles["Body10_H3"]) for x in r.methods_procedure], bulletType="bullet"))
            if include["t_trial"].get():
                story.append(Paragraph("2.3.4. Trial History", styles["H3_10"]))
                story.append(Spacer(1, 6))
                if r.trial_history:
                    data_th = [["Trial#", "Issue", "Possible Reasons"]] + [[t.number, t.issue, t.reasons] for t in r.trial_history]
                    tt = Table(data_th, colWidths=[1.0*inch, 2.0*inch, 4.0*inch])
                    tt.setStyle(TableStyle([
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]))
                    story.append(tt)

        if include["t_results"].get():
            story.append(Paragraph("2.4. Results", styles["H2_11"]))
            for i, it in enumerate(r.results, 1):
                if it.kind == "text":
                    story.append(Paragraph(f"2.4.{i}. {it.title}", styles["H3_10"]))
                    story.append(Paragraph((it.content or "").replace("\n", "<br/>"), styles["Body10_H3"]))
                elif it.kind == "table":
                    story.append(Paragraph(f"{it.title}", styles["H3_10"]))
                    data = it.table_data if it.table_data else parse_table_text(it.content)
                    if data:
                        ncols = max(len(rw) for rw in data)
                        total_w = 7.0 * inch
                        col_w = [total_w / max(1, ncols)] * ncols
                        tbl = Table(data, colWidths=col_w)
                        tbl.setStyle(TableStyle([
                            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ]))
                        story.append(tbl)
                    else:
                        story.append(Paragraph((it.content or "").replace("\n", "<br/>"), styles["Body10_H3"]))
                elif it.kind == "image":
                    imgs = it.images if it.images else ([it.image_path] if it.image_path else [])
                    if imgs:
                        if it.title: story.append(Paragraph(f"{it.title}", styles["H3_10"]))
                        for pth in imgs:
                            if pth and os.path.isfile(pth):
                                story.append(RLImage(pth, width=5.8 * inch, height=3.3 * inch))
                                if it.caption:
                                    story.append(Paragraph(it.caption, styles["Body10_H3"]))

        if include["t_conc"].get():
            story.append(Paragraph("2.5. Conclusion", styles["H2_11"]))
            story.append(ListFlowable([Paragraph(x, styles["Body10_H2"]) for x in r.conclusion], bulletType="bullet"))

        if include["t_misc"].get():
            story.append(Paragraph("2.6. Miscellaneous", styles["H2_11"]))
            story.append(Paragraph((r.miscellaneous or "").replace("\n", "<br/>"), styles["Body10_H2"]))

        if include["t_refs"].get():
            story.append(Paragraph("2.7. References", styles["H2_11"]))
            story.append(ListFlowable([Paragraph(x, styles["Body10_H2"]) for x in r.references], bulletType="bullet"))

        if include["sec_reg"].get():
            story.append(Paragraph("3. Regulatory", styles["H1_12"]))
            story.append(Paragraph("3.1. Application regulations", styles["H2_11"]))
            story.append(Paragraph((r.regulations or "").replace("\n", "<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("3.2. Label Requirements", styles["H2_11"]))
            story.append(Paragraph((r.label_req or "").replace("\n", "<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("3.3. Certification Requirements", styles["H2_11"]))
            story.append(Paragraph((r.certification_req or "").replace("\n", "<br/>"), styles["Body10_H2"]))

        # 4. Scale Up
        if include["sec_scale"].get():
            story.append(Paragraph("4. Scale Up", styles["H1_12"]))

            story.append(Paragraph("4.1. Manufacturing Order", styles["H2_11"]))
            if r.manuf_order_steps:
                story.append(ListFlowable([Paragraph(x, styles["Body10_H2"]) for x in r.manuf_order_steps],
                                          bulletType="bullet"))

            story.append(Paragraph("4.2. Formulation Risk", styles["H2_11"]))
            story.append(Paragraph((r.formulation_risk_text or "").replace("\n","<br/>"), styles["Body10_H2"]))

            story.append(Paragraph("4.3. Hazards", styles["H2_11"]))
            story.append(Paragraph((r.hazards_text or "").replace("\n","<br/>"), styles["Body10_H2"]))

            story.append(Paragraph("4.4. Equipment", styles["H2_11"]))
            story.append(Paragraph((r.equipment_text or "").replace("\n","<br/>"), styles["Body10_H2"]))

            story.append(Paragraph("4.5. CAPEX Requirements", styles["H2_11"]))
            story.append(Paragraph((r.capex_text or "").replace("\n","<br/>"), styles["Body10_H2"]))

            story.append(Paragraph("4.6. Safety Assessment", styles["H2_11"]))
            story.append(Paragraph((r.safety_assess_text or "").replace("\n","<br/>"), styles["Body10_H2"]))

        if include["sec_quality"].get():
            story.append(Paragraph("5. Quality", styles["H1_12"]))
            story.append(Paragraph("5.1. Raw Material Sourcing", styles["H2_11"]))
            story.append(Paragraph((r.raw_material_sourcing or "").replace("\n", "<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("5.2. LIMS Setup", styles["H2_11"]))
            story.append(Paragraph((r.lims_setup or "").replace("\n", "<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("5.3. Stability Testing", styles["H2_11"]))
            story.append(Paragraph((r.stability_testing or "").replace("\n", "<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("5.4. Packaging Compatibility", styles["H2_11"]))
            story.append(Paragraph((r.packaging_compatibility or "").replace("\n", "<br/>"), styles["Body10_H2"]))

        if include["sec_commercial"].get():
            story.append(Paragraph("6. Commercial", styles["H1_12"]))
            story.append(Paragraph("6.1. Customer Objectives / Problem Statement", styles["H2_11"]))
            story.append(Paragraph((r.c_obj_problem or "").replace("\n","<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("6.2. SMART Success Criteria", styles["H2_11"]))
            for key,label in [("S","Specific"),("M","Measurable"),("A","Achievable"),("R","Relevant"),("T","Time-bound")]:
                val = r.smart_goals.get(key,"").strip()
                if val:
                    story.append(Paragraph(f"{label}: {val}", styles["Body10_H3"]))
            story.append(Paragraph("6.3. Customer Specifications", styles["H2_11"]))
            story.append(Paragraph((r.c_specs or "").replace("\n","<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("6.4. Expected Business Volume", styles["H2_11"]))
            story.append(Paragraph((r.c_expected_volume or "").replace("\n","<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("6.5. Packaging Requirement", styles["H2_11"]))
            story.append(Paragraph((r.c_packaging_req or "").replace("\n","<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("6.6. Raw Material Restrictions / Preferences", styles["H2_11"]))
            story.append(Paragraph((r.c_raw_material_prefs or "").replace("\n","<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("6.7. Sample Needed", styles["H2_11"]))
            story.append(Paragraph((r.c_sample_needed or "").replace("\n","<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("6.8. Opportunity Timeline", styles["H2_11"]))
            story.append(Paragraph((r.c_opportunity_timeline or "").replace("\n","<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("6.9. Target Application", styles["H2_11"]))
            story.append(Paragraph((r.target_application or "").replace("\n","<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("6.10. Customer Feedback", styles["H2_11"]))
            story.append(Paragraph((r.customer_feedback or "").replace("\n","<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("6.11. TDS Development", styles["H2_11"]))
            story.append(Paragraph((r.tds_development or "").replace("\n","<br/>"), styles["Body10_H2"]))
            story.append(Paragraph("6.12. Email Correspondence", styles["H2_11"]))
            if r.email_correspondence:
                data = [["Date","Customer Name","Correspondence"]] + [[e.date, e.customer, e.correspondence] for e in r.email_correspondence]
                tt = Table(data, colWidths=[1.3*inch, 2.0*inch, 3.7*inch])
                tt.setStyle(TableStyle([
                    ("GRID", (0,0), (-1,-1), 0.25, colors.grey),
                    ("BACKGROUND", (0,0), (-1,0), colors.whitesmoke),
                    ("VALIGN", (0,0), (-1,-1), "TOP"),
                ]))
                story.append(tt)

        def on_page(canvas, doc_obj):
            # Header band at the very top region
            if header_img and os.path.isfile(header_img):
                try:
                    canvas.drawImage(
                        header_img,
                        x=doc_obj.leftMargin,
                        y=doc_obj.pagesize[1] - 0.85 * inch,  # near top edge
                        width=doc_obj.width,
                        height=0.5 * inch,
                        preserveAspectRatio=True,
                        mask='auto'
                    )
                except Exception:
                    pass

            # Footer band near bottom + page number
            if footer_img and os.path.isfile(footer_img):
                try:
                    canvas.drawImage(
                        footer_img,
                        x=doc_obj.leftMargin,
                        y=0.35 * inch,
                        width=doc_obj.width,
                        height=0.5 * inch,
                        preserveAspectRatio=True,
                        mask='auto'
                    )
                except Exception:
                    pass

            page_num = canvas.getPageNumber()
            canvas.setFont("Helvetica", 9)
            canvas.drawRightString(doc_obj.pagesize[0]-doc_obj.rightMargin, 0.25 * inch, f"{page_num}")

        # Leave space for header/footer bands
        doc = SimpleDocTemplate(
            out_path,
            pagesize=A4,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=1.35 * inch,     # room for header band
            bottomMargin=1.2 * inch    # room for footer band
        )
        doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
        messagebox.showinfo("PDF saved", f"Saved: {out_path}")

    def _generate(self):
        r = self._collect()
        if not r.project_title:
            messagebox.showwarning("Missing title", "Please enter a Project Title in General Information.")
            return

        out_dir = self.out_dir_var.get().strip()
        ensure_dir(os.path.join(out_dir, "_"))
        base = f"{r.project_title.strip().replace(os.sep, '_')}_{datetime.date.today().isoformat()}"
        docx_path = os.path.join(out_dir, base + ".docx")
        pdf_path = os.path.join(out_dir, base + ".pdf")
        want = self.export_fmt.get()
        header_img = self.header_path_var.get().strip() if os.path.isfile(self.header_path_var.get().strip()) else None
        footer_img = self.footer_path_var.get().strip() if os.path.isfile(self.footer_path_var.get().strip()) else None

        if want in ("docx", "both"):
            if Document is None:
                messagebox.showerror("python-docx not found", "Install python-docx to export .docx")
                return
            self._export_docx(r, docx_path, header_img, footer_img)
            messagebox.showinfo("DOCX saved", f"Saved: {docx_path}")

        if want in ("pdf", "both"):
            if docx2pdf_convert and os.path.isfile(docx_path):
                try:
                    docx2pdf_convert(docx_path, pdf_path)
                    messagebox.showinfo("PDF saved", f"Saved: {pdf_path}")
                except Exception:
                    self._export_pdf_basic(r, pdf_path, header_img, footer_img)
            else:
                self._export_pdf_basic(r, pdf_path, header_img, footer_img)

# ------------------------------ Entrypoint -----------------------------------
if __name__ == "__main__":
    app = ReportApp()
    app.mainloop()
